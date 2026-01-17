import os
import uuid
import json
import json as json_module
import requests
import logging
import asyncio
from datetime import datetime, timedelta
from typing import List, Optional
from contextlib import asynccontextmanager
from fastapi import FastAPI, File, UploadFile, Depends, HTTPException, BackgroundTasks, Form
from fastapi.middleware.cors import CORSMiddleware
from middleware import (
    ErrorHandlingMiddleware, 
    SecurityHeadersMiddleware, 
    RequestValidationMiddleware, 
    RequestLoggingMiddleware,
    BusinessLogicError,
    AIProcessingError,
    FileProcessingError
)
from sqlalchemy.orm import Session
from sqlalchemy import or_
from database import get_db
from models import Document, Org, User, Framework, Control, Requirement, EvidenceLink, Scan, ScanResult, Gap, DocumentControlLink, DocumentPage, Settings
from storage import storage
from worker_tasks import extract_document_text, process_scan
from pydantic import BaseModel
from init_db import initialize_database

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# NOT A SECRET: This is a placeholder key for local AI endpoints (Ollama, LM Studio, etc.)
# that don't require authentication. It has no security value and is never used with real APIs.
# The OpenAI SDK requires an api_key parameter, so we provide this dummy value for local endpoints.
LOCAL_AI_PLACEHOLDER_KEY = os.getenv('LOCAL_AI_PLACEHOLDER_KEY', 'sk-local-endpoint-no-auth')

def _analyze_document_two_step(file_text: str, filename: str, available_controls: List[dict], ai_client) -> List[dict]:
    """
    Two-step document analysis:
    1. First scan and summarize the document
    2. Then map the summary to compliance controls
    """
    if isinstance(ai_client, dict) and ai_client.get('type') == 'ollama':
        return _analyze_document_ollama_two_step(file_text, filename, available_controls, ai_client)
    else:
        return _analyze_document_openai_two_step(file_text, filename, available_controls, ai_client)

def _analyze_document_ollama_two_step(file_text: str, filename: str, available_controls: List[dict], ai_client: dict) -> List[dict]:
    """JSON-to-JSON two-step analysis specifically for Ollama models"""
    import requests
    import json as json_module
    
    endpoint = ai_client['endpoint']
    model = ai_client['model']
    
    # Step 1: Document Scanning - Create JSON Summary
    scan_prompt = f"""Analyze document: {filename}
Content: {file_text[:5000]}

Return JSON summary:
{{"document_type":"screenshot","primary_topic":"main subject","key_content_indicators":["keywords found"],"security_areas":["security domain"],"main_requirements":["core requirement"],"distinguishing_features":"what makes this unique"}}"""
    
    logger.info(f"Step 1: Creating JSON summary for {filename}")
    
    try:
        # Use only generate API for completions
        scan_response = requests.post(
            f"{endpoint}/api/generate",
            json={
                "model": model,
                "prompt": scan_prompt,
                "stream": False,
                "options": {
                    "temperature": 0.1,
                    "num_predict": 1000,  # Increased for complete JSON responses
                    "num_ctx": int(os.getenv("OLLAMA_CONTEXT_SIZE", "32768")),
                    "stop": ["\n\n\n"]  # Only stop on triple newlines
                }
            },
            timeout=90
        )
        
        if scan_response.status_code != 200:
            logger.error(f"Step 1 failed for {filename}: {scan_response.status_code}")
            return generate_fallback_suggestions_from_filename(filename, available_controls)
        
        scan_result = scan_response.json()
        
        # Use only generate API response format
        document_summary_raw = scan_result.get('response', '')
        logger.info(f"Step 1: Using generate API response for {filename}")
        
        # NEVER use thinking field - log but ignore
        if 'thinking' in scan_result and scan_result.get('thinking'):
            logger.info(f"Step 1: Thinking field ignored for {filename}: {scan_result.get('thinking', '')[:100]}...")
        
        if not document_summary_raw:
            logger.warning(f"Step 1 produced empty summary for {filename}")
            return generate_fallback_suggestions_from_filename(filename, available_controls)
        
        logger.info(f"Step 1 raw response for {filename}: {document_summary_raw[:300]}...")
        
        # Parse the JSON summary from Step 1
        document_summary_json = _extract_json_from_response(document_summary_raw)
        if not document_summary_json:
            logger.warning(f"Step 1 failed to produce valid JSON for {filename}")
            return generate_fallback_suggestions_from_filename(filename, available_controls)
        
        logger.info(f"Step 1 JSON summary for {filename}: {document_summary_json}")
        
        # Step 2: Control Mapping using the JSON from Step 1
        controls_json = []
        for i, control in enumerate(available_controls[:20], 1):
            controls_json.append({
                "number": i,
                "code": control['code'],
                "title": control['title'],
                "framework": control.get('framework', 'Unknown')
            })
        
        mapping_prompt = f"""Document Summary: {json_module.dumps(document_summary_json, indent=2)}

Available Controls: {json_module.dumps(controls_json, indent=2)}

Return control mapping JSON:
{{"selected_control_number":1,"confidence":0.90,"reasoning":"Brief match explanation"}}"""
        
        logger.info(f"Step 2: JSON mapping controls for {filename}")
        
        # Use only generate API for completions
        mapping_response = requests.post(
            f"{endpoint}/api/generate",
            json={
                "model": model,
                "prompt": mapping_prompt,
                "stream": False,
                "options": {
                    "temperature": 0.1,
                    "num_predict": 600,  # Increased for complete JSON responses with reasoning
                    "num_ctx": int(os.getenv("OLLAMA_CONTEXT_SIZE", "32768")),
                    "stop": ["\n\n\n"]  # Only stop on triple newlines
                }
            },
            timeout=90
        )
        
        if mapping_response.status_code != 200:
            logger.error(f"Step 2 failed for {filename}: {mapping_response.status_code}")
            return generate_fallback_suggestions_from_filename(filename, available_controls)
        
        mapping_result = mapping_response.json()
        
        # Use only generate API response format
        mapping_text_raw = mapping_result.get('response', '')
        logger.info(f"Step 2: Using generate API response for {filename}")
        
        # NEVER use thinking field - log but ignore
        if 'thinking' in mapping_result and mapping_result.get('thinking'):
            logger.info(f"Step 2: Thinking field ignored for {filename}: {mapping_result.get('thinking', '')[:100]}...")
        
        if not mapping_text_raw:
            logger.warning(f"Step 2 produced empty mapping for {filename}")
            return generate_fallback_suggestions_from_filename(filename, available_controls)
        
        logger.info(f"Step 2 raw response for {filename}: {mapping_text_raw}")
        
        # Parse the JSON mapping result from Step 2
        mapping_json = _extract_json_from_response(mapping_text_raw)
        if not mapping_json:
            logger.warning(f"Step 2 failed to produce valid JSON for {filename}")
            return generate_fallback_suggestions_from_filename(filename, available_controls)
        
        logger.info(f"Step 2 JSON mapping for {filename}: {mapping_json}")
        
        # Convert the JSON result to the expected format
        return _convert_json_mapping_to_suggestions(mapping_json, available_controls, controls_json)
        
    except Exception as e:
        logger.error(f"JSON-to-JSON two-step analysis failed for {filename}: {e}")
        return generate_fallback_suggestions_from_filename(filename, available_controls)

def _convert_json_mapping_to_suggestions(mapping_json: dict, available_controls: List[dict], controls_json: List[dict]) -> List[dict]:
    """Convert JSON mapping result to the expected suggestions format"""
    try:
        selected_number = mapping_json.get('selected_control_number', 1)
        confidence = mapping_json.get('confidence', 0.5)
        reasoning = mapping_json.get('reasoning', 'AI analysis')
        
        # Validate the selected control number
        if 1 <= selected_number <= len(controls_json):
            control_index = selected_number - 1  # Convert to 0-indexed
            if control_index < len(available_controls):
                control = available_controls[control_index]
                return [{
                    'control_code': control['code'],
                    'control_title': control['title'],
                    'framework_name': control.get('framework', 'Unknown'),
                    'confidence': max(0.0, min(1.0, float(confidence))),
                    'reasoning': reasoning[:200]  # Limit length
                }]
        
        logger.warning(f"Invalid control number {selected_number} in JSON mapping")
        return []
        
    except Exception as e:
        logger.error(f"Failed to convert JSON mapping: {e}")
        return []

def _parse_structured_control_response(response_text: str, available_controls: List[dict]) -> List[dict]:
    """Parse the structured NUMBER/REASONING response format"""
    import re
    
    try:
        # Extract number and reasoning
        number_match = re.search(r'NUMBER:\s*(\d+)', response_text, re.IGNORECASE)
        reasoning_match = re.search(r'REASONING:\s*(.+)', response_text, re.IGNORECASE | re.DOTALL)
        
        if not number_match:
            logger.warning("Could not find NUMBER in structured response")
            return []
        
        control_number = int(number_match.group(1)) - 1  # Convert to 0-indexed
        reasoning = reasoning_match.group(1).strip() if reasoning_match else "AI analysis"
        
        # Validate control number
        if 0 <= control_number < len(available_controls):
            control = available_controls[control_number]
            return [{
                'control_code': control['code'],
                'control_title': control['title'],
                'framework_name': control.get('framework', 'Unknown'),
                'confidence': 0.85,  # High confidence for structured response
                'reasoning': reasoning[:200]  # Limit length
            }]
        else:
            logger.warning(f"Invalid control number: {control_number + 1}")
            return []
            
    except Exception as e:
        logger.error(f"Failed to parse structured response: {e}")
        return []

def _analyze_document_openai_two_step(file_text: str, filename: str, available_controls: List[dict], ai_client) -> List[dict]:
    """Two-step analysis for OpenAI (can use structured outputs natively)"""
    # For now, use existing OpenAI logic but could be enhanced with structured outputs
    return []

def generate_fallback_suggestions_from_filename(filename: str, available_controls: List[dict]) -> List[dict]:
    """Generate control suggestions based on filename when AI analysis fails."""
    filename_lower = filename.lower()
    
    # Common patterns for compliance documents
    suggestions = []
    
    for control in available_controls[:10]:  # Check first 10 controls
        confidence = 0.3  # Default low confidence for filename matching
        reasoning = f"Filename-based suggestion for {filename}"
        
        # Patch management documents
        if any(word in filename_lower for word in ['patch', 'update', 'os']):
            if any(word in control['code'].lower() or word in control['title'].lower() 
                   for word in ['patch', 'update', 'os', 'operating']):
                confidence = 0.7
                reasoning = f"Document name suggests patch management, matching {control['code']}"
        
        # Access control documents
        elif any(word in filename_lower for word in ['access', 'auth', 'mfa', 'login']):
            if any(word in control['code'].lower() or word in control['title'].lower() 
                   for word in ['access', 'auth', 'mfa', 'authentication']):
                confidence = 0.7
                reasoning = f"Document name suggests access control, matching {control['code']}"
        
        # Backup documents
        elif any(word in filename_lower for word in ['backup', 'recovery']):
            if any(word in control['code'].lower() or word in control['title'].lower() 
                   for word in ['backup', 'recovery']):
                confidence = 0.7
                reasoning = f"Document name suggests backup/recovery, matching {control['code']}"
        
        # Application control documents
        elif any(word in filename_lower for word in ['app', 'software', 'application']):
            if any(word in control['code'].lower() or word in control['title'].lower()
                   for word in ['application', 'software']):
                confidence = 0.4  # Lower confidence for filename-only matching
                reasoning = f"Document name suggests application control, matching {control['code']}"

        if confidence > 0.3:  # Raised threshold for suggestions
            suggestions.append({
                'control_code': control['code'],
                'control_title': control['title'],
                'framework_name': control.get('framework', 'Unknown'),
                'confidence': confidence,
                'reasoning': reasoning
            })
    
    # Return top 3 suggestions
    return sorted(suggestions, key=lambda x: x['confidence'], reverse=True)[:3]

def _extract_json_content_only(response_text):
    """Extract just the JSON object as a string, no parsing."""
    if not response_text:
        return None

    import re

    # Clean the response first
    response_text = response_text.strip()

    # If it already looks like valid JSON, return it
    if response_text.startswith('{') and response_text.endswith('}'):
        return response_text

    # Try to find the outermost JSON object by counting braces
    start_idx = response_text.find('{')
    if start_idx == -1:
        return None

    brace_count = 0
    end_idx = start_idx

    for i in range(start_idx, len(response_text)):
        if response_text[i] == '{':
            brace_count += 1
        elif response_text[i] == '}':
            brace_count -= 1
            if brace_count == 0:
                end_idx = i + 1
                break

    if brace_count == 0 and end_idx > start_idx:
        return response_text[start_idx:end_idx].strip()

    return None

def _extract_json_from_response(response_text):
    """Extract JSON from a response that might contain extra text."""
    if not response_text:
        logger.warning("Empty response text for JSON extraction")
        return None
    
    # Try to find JSON object in the response
    import re
    import json
    
    # Clean the response text
    response_text = response_text.strip()
    logger.info(f"Extracting JSON from response (length: {len(response_text)}): {response_text[:200]}...")
    
    # First, try to parse the entire response as JSON
    try:
        parsed = json.loads(response_text)
        if isinstance(parsed, dict):
            logger.info("Successfully parsed entire response as JSON")
            return parsed
    except json.JSONDecodeError:
        pass
    
    # Look for JSON object patterns (more comprehensive)
    json_patterns = [
        r'\{[^{}]*"document_type"[^{}]*\}',  # Document summary JSON
        r'\{[^{}]*"selected_control_number"[^{}]*\}',  # Control mapping JSON
        r'\{[^{}]*"suggestions"[^{}]*\[[^\]]*\][^{}]*\}',  # Suggestions JSON
        r'\{.*?"suggestions".*?\[.*?\].*?\}',  # More flexible suggestions pattern
        r'\{[\s\S]*?\}',  # Any JSON object (non-greedy)
    ]
    
    for i, pattern in enumerate(json_patterns):
        matches = re.findall(pattern, response_text, re.DOTALL)
        for match in matches:
            try:
                # Clean the match
                cleaned_match = match.strip()
                parsed = json.loads(cleaned_match)
                if isinstance(parsed, dict):
                    logger.info(f"Successfully extracted JSON using pattern {i}: {cleaned_match[:100]}...")
                    return parsed
            except json.JSONDecodeError:
                continue
    
    # Try to find JSON between code block markers
    code_block_patterns = [
        r'```json\s*(\{.*?\})\s*```',
        r'```\s*(\{.*?\})\s*```',
    ]
    
    for pattern in code_block_patterns:
        matches = re.findall(pattern, response_text, re.DOTALL)
        for match in matches:
            try:
                parsed = json.loads(match.strip())
                if isinstance(parsed, dict):
                    return parsed
            except json.JSONDecodeError:
                continue
    
    # Try the whole response as-is
    try:
        return json.loads(response_text)
    except json.JSONDecodeError:
        pass
    
    # Last resort: try to extract just the JSON portion
    try:
        # Find first { and last }
        start = response_text.find('{')
        end = response_text.rfind('}')
        if start >= 0 and end > start:
            json_portion = response_text[start:end+1]
            return json.loads(json_portion)
    except (json.JSONDecodeError, ValueError):
        pass
    
    return None

def _safe_json_loads(json_data, default=None):
    """Safely parse JSON data from JSONB columns, which are already parsed by PostgreSQL."""
    if json_data is None:
        return default

    # JSONB columns return already-parsed Python objects (str, list, dict)
    # Just return them directly - no JSON parsing needed
    if isinstance(json_data, (str, list, dict)):
        return json_data if json_data else default

    # For any other type, log and return default
    logger.warning(f"Unexpected data type in database - returning default. Type: {type(json_data)}, Content: {str(json_data)[:100]}...")
    return default

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    logger.info("Starting GeekyGoose Compliance API...")
    
    # Initialize database on startup
    if not initialize_database():
        logger.error("Database initialization failed!")
        raise RuntimeError("Database initialization failed")
    
    # Start the periodic AI retry task
    retry_task = asyncio.create_task(periodic_ai_retry_task())
    logger.info("Started periodic AI retry task")
    
    logger.info("GeekyGoose Compliance API startup complete")
    yield
    # Shutdown
    logger.info("GeekyGoose Compliance API shutting down...")
    retry_task.cancel()
    try:
        await retry_task
    except asyncio.CancelledError:
        logger.info("Periodic AI retry task cancelled")

app = FastAPI(
    title="GeekyGoose Compliance API",
    description="Compliance automation platform for SMB + internal IT teams",
    version="0.3.0",  # Updated version
    lifespan=lifespan
)

# Add security and error handling middleware (order matters!)
app.add_middleware(ErrorHandlingMiddleware)
app.add_middleware(SecurityHeadersMiddleware)
app.add_middleware(RequestValidationMiddleware)
app.add_middleware(RequestLoggingMiddleware)

# CORS middleware - Allow requests from Next.js container and localhost
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000", 
        "http://127.0.0.1:3000",
        "http://web:3000",  # Allow from Next.js container
        "*"  # Allow all origins since Next.js will proxy requests
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Allowed file types
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "image/png",
    "image/jpeg"
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

async def analyze_file_content_for_controls(file: UploadFile, file_content: bytes) -> list:
    """Analyze file content and suggest relevant compliance controls."""
    try:
        # Get available templates/controls from the database
        from database import SessionLocal
        db = SessionLocal()
        
        try:
            # Get all controls from database
            controls = db.query(Control).all()
            logger.info(f"Found {len(controls)} controls in database for analysis")
            
            if not controls:
                logger.warning("No controls found in database - cannot provide AI analysis")
                return []
            
            # Convert to the format expected by our analysis
            available_controls = []
            for control in controls:
                available_controls.append({
                    'code': control.code,
                    'title': control.title,
                    'framework': getattr(control.framework, 'name', 'Unknown') if hasattr(control, 'framework') else 'Unknown',
                    'description': control.description or '',
                })
            
            logger.info(f"Prepared {len(available_controls)} controls for AI analysis")
        finally:
            db.close()
        
        if not available_controls:
            return []
        
        # Extract content based on file type with enhanced detection
        file_text = ""
        filename = file.filename.split('\\')[-1].split('/')[-1]  # Clean filename
        
        # Use python-magic for better file type detection
        try:
            import magic
            file_mime = magic.from_buffer(file_content, mime=True)
        except (ImportError, Exception) as e:
            logger.debug(f"Magic library not available or failed ({e}), using file content_type")
            file_mime = getattr(file, 'content_type', 'text/plain')
        
        file_content_type = getattr(file, 'content_type', 'text/plain')
        if file_mime == "text/plain" or file_content_type == "text/plain" or filename.lower().endswith(('.txt', '.md', '.csv')):
            try:
                import chardet
                # Detect encoding for better text extraction
                detected = chardet.detect(file_content)
                encoding = detected.get('encoding', 'utf-8')
                file_text = file_content.decode(encoding, errors='ignore')[:2000]  # Limit to first 2000 chars
            except Exception as e:
                try:
                    file_text = file_content.decode('utf-8', errors='ignore')[:2000]
                except:
                    file_text = f"Text file: {filename} (encoding issue)"
                
        elif file_mime == "application/pdf" or file.content_type == "application/pdf" or filename.lower().endswith('.pdf'):
            try:
                import fitz  # PyMuPDF
                import pdfplumber
                
                # Try PyMuPDF first (faster)
                try:
                    pdf_doc = fitz.open(stream=file_content, filetype="pdf")
                    text_pages = []
                    for page_num in range(min(3, pdf_doc.page_count)):  # First 3 pages
                        page = pdf_doc[page_num]
                        page_text = page.get_text()
                        if page_text.strip():
                            text_pages.append(page_text[:800])  # More text per page
                    
                    pdf_doc.close()
                    file_text = " ".join(text_pages)
                except Exception:
                    # Fallback to pdfplumber for more complex PDFs
                    import io
                    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                        text_pages = []
                        for i, page in enumerate(pdf.pages[:3]):
                            page_text = page.extract_text()
                            if page_text:
                                text_pages.append(page_text[:800])
                        file_text = " ".join(text_pages)
                
                if not file_text.strip():
                    file_text = f"PDF document: {filename} (text extraction failed)"
            except Exception as e:
                logger.warning(f"PDF extraction failed for {filename}: {e}")
                file_text = f"PDF document: {filename} (text extraction failed)"
                
        elif (file_mime and file_mime.startswith("image/")) or file.content_type.startswith("image/") or filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
            # Enhanced image processing with OCR fallback
            try:
                from ai_scanner import get_ai_client
                import base64
                from PIL import Image
                import io
                
                # Process image with PIL first
                try:
                    image = Image.open(io.BytesIO(file_content))
                    # Convert to RGB if needed for better AI analysis
                    if image.mode != 'RGB':
                        image = image.convert('RGB')
                    
                    # Resize if too large (for better AI processing)
                    max_size = (1024, 1024)
                    if image.size[0] > max_size[0] or image.size[1] > max_size[1]:
                        image.thumbnail(max_size, Image.Resampling.LANCZOS)
                        
                        # Convert back to bytes
                        img_buffer = io.BytesIO()
                        image.save(img_buffer, format='JPEG', quality=85)
                        processed_content = img_buffer.getvalue()
                    else:
                        processed_content = file_content
                        
                except Exception:
                    processed_content = file_content
                
                ai_client = get_ai_client()
                image_b64 = base64.b64encode(processed_content).decode('utf-8')
                
                if not isinstance(ai_client, dict):  # OpenAI
                    try:
                        response = ai_client.chat.completions.create(
                            model="gpt-4o",  # Updated to latest vision model
                            messages=[
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": f"Analyze this image thoroughly and describe any visible text, error messages, security configurations, system interfaces, compliance-related information, policies, procedures, or other relevant content you can see. Focus on compliance and security aspects. Image: {filename}"
                                        },
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:image/jpeg;base64,{image_b64}",
                                                "detail": "high"  # High detail for better text recognition
                                            }
                                        }
                                    ]
                                }
                            ],
                            max_tokens=500
                        )
                        file_text = f"Image analysis: {response.choices[0].message.content}"
                    except Exception as e:
                        logger.warning(f"Vision AI failed for {filename}: {e}")
                        # Fallback to OCR
                        try:
                            import pytesseract
                            from PIL import Image
                            image = Image.open(io.BytesIO(file_content))
                            ocr_text = pytesseract.image_to_string(image)
                            if ocr_text.strip():
                                file_text = f"OCR extracted text from {filename}: {ocr_text[:500]}"
                            else:
                                file_text = f"Screenshot/Image: {filename} (no text detected)"
                        except Exception:
                            file_text = f"Screenshot/Image: {filename}"
                else:
                    # Try OCR for Ollama users
                    try:
                        import pytesseract
                        from PIL import Image
                        image = Image.open(io.BytesIO(file_content))
                        ocr_text = pytesseract.image_to_string(image)
                        if ocr_text.strip():
                            file_text = f"OCR extracted text from {filename}: {ocr_text[:500]}"
                        else:
                            file_text = f"Screenshot/Image: {filename}"
                    except Exception:
                        file_text = f"Image: {filename}"
            except Exception as e:
                logger.error(f"Image processing failed for {filename}: {e}")
                file_text = f"Image: {filename}"
                
        elif file_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or filename.lower().endswith('.docx'):
            # Enhanced Word document processing
            try:
                from docx import Document
                import io
                
                doc = Document(io.BytesIO(file_content))
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text)
                
                file_text = "\n".join(paragraphs[:30])  # First 30 paragraphs/cells
                if not file_text.strip():
                    file_text = f"Word document: {filename} (text extraction failed)"
            except Exception as e:
                logger.warning(f"Word document processing failed for {filename}: {e}")
                file_text = f"Word document: {filename} (text extraction failed)"
                
        elif filename.lower().endswith(('.html', '.htm')):
            # HTML document processing
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(file_content, 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                file_text = soup.get_text()[:2000]
            except Exception:
                file_text = f"HTML document: {filename}"
                
        elif filename.lower().endswith('.md'):
            # Markdown document processing
            try:
                import markdown
                md_text = file_content.decode('utf-8', errors='ignore')
                html = markdown.markdown(md_text)
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html, 'html.parser')
                file_text = soup.get_text()[:2000]
            except Exception:
                try:
                    file_text = file_content.decode('utf-8', errors='ignore')[:2000]
                except:
                    file_text = f"Markdown document: {filename}"
        else:
            file_text = f"Document: {filename} (type: {file_mime or file.content_type})"
        
        # Create analysis prompt
        controls_context = "\n".join([
            f"- {c['code']}: {c['title']} ({c['framework']}) - {c['description'][:500]}..."
            for c in available_controls[:50]  # Increased from 10 to 50 controls, longer descriptions
        ])
        
        analysis_prompt = f"""
You are a compliance expert. Analyze this document and identify which compliance controls it relates to.

Document: {filename}
Content: {file_text[:5000]}{'...' if len(file_text) > 50000 else ''}

Available compliance controls:
{controls_context}

⚠️ CRITICAL MAPPING RULES - Follow these EXACTLY:

1. **MACRO/OFFICE SECURITY** = EE-3 ONLY:
   - If you see: Microsoft Office, Trust Center, macro settings, VBA, macro security, Excel/Word security
   - Then MUST use: EE-3 "Configure Microsoft Office Macro Settings"
   - DO NOT use EE-1 or EE-4 for macro content!

2. **OS UPDATES/PATCHES** = EE-6:
   - Windows Update, OS patches, system updates, firmware updates
   - Use: EE-6 "Patch Operating Systems"

3. **APP UPDATES** = EE-2:
   - Application updates, software patches (not OS)
   - Use: EE-2 "Patch Applications"

4. **AUTHENTICATION** = EE-5:
   - MFA, 2FA, multi-factor authentication
   - Use: EE-5 "Multi-Factor Authentication"

5. **BACKUPS** = EE-7:
   - Backup settings, recovery points
   - Use: EE-7 "Backup Data"

6. **APP WHITELISTING** = EE-1:
   - Application whitelisting, AppLocker, execution control
   - Use: EE-1 "Application Control"
   - NOT for macro settings!

7. **BROWSER HARDENING** = EE-4:
   - Browser security, application hardening
   - Use: EE-4 "User Application Hardening"

8. **ADMIN ACCESS** = EE-8:
   - Admin restrictions, privileged access
   - Use: EE-8 "Restrict Admin Privileges"

Analyze the document content and provide the top 3 most relevant controls.
For each control, provide:
- control_code: The exact control code
- control_title: The exact control title
- framework_name: The framework name
- confidence: A number between 0.0 and 1.0
- reasoning: Brief explanation (1-2 sentences)

Respond ONLY with valid JSON in this exact format:
{{
  "suggestions": [
    {{
      "control_code": "CONTROL_CODE",
      "control_title": "Control Title",
      "framework_name": "Framework Name",
      "confidence": 0.8,
      "reasoning": "Brief explanation of why this control is relevant."
    }}
  ]
}}

Do not include any text before or after the JSON. Return empty suggestions array if no relevant controls found."""
        
        # Call AI for analysis using new two-step approach
        try:
            from ai_scanner import get_ai_client
            logger.info(f"Attempting to get AI client for {filename}")
            ai_client = get_ai_client()
            logger.info(f"AI client initialized: {type(ai_client)}")
            
            # Use the new two-step analysis approach
            logger.info(f"Starting two-step analysis for {filename}")
            suggested_controls = _analyze_document_two_step(
                file_text=file_text,
                filename=filename,
                available_controls=available_controls,
                ai_client=ai_client
            )
            
            if suggested_controls:
                logger.info(f"Two-step analysis succeeded for {filename}, got {len(suggested_controls)} suggestions")
                return suggested_controls
            else:
                logger.warning(f"Two-step analysis failed for {filename}, falling back to original method")
            
            # Fallback to original method if two-step fails
            if isinstance(ai_client, dict) and ai_client.get('type') == 'ollama':
                import requests
                endpoint = ai_client['endpoint']
                model = ai_client['model']
                logger.info(f"Using Ollama at {endpoint} with model {model}")
                
                # Test Ollama connectivity first
                try:
                    test_response = requests.get(f"{endpoint}/api/tags", timeout=10)
                    logger.info(f"Ollama connectivity test: {test_response.status_code}")
                    if test_response.status_code != 200:
                        logger.error(f"Ollama not reachable at {endpoint}")
                        return generate_fallback_suggestions_from_filename(filename, available_controls)
                except Exception as conn_error:
                    logger.error(f"Cannot connect to Ollama at {endpoint}: {conn_error}")
                    return generate_fallback_suggestions_from_filename(filename, available_controls)
                
                # Clear and simple prompt
                simple_prompt = f"""Analyze document: {filename}
Content: {file_text[:5000] if file_text else 'Filename analysis only'}

Available controls: {[c['code'] for c in available_controls[:3]]}

Respond with JSON only:
{{"suggestions":[{{"control_code":"{available_controls[0]['code'] if available_controls else 'EE-1'}","control_title":"Title","framework_name":"Essential Eight","confidence":0.8,"reasoning":"Why this matches"}}]}}"""
                
                logger.info(f"Sending prompt to Ollama (length: {len(simple_prompt)})")
                
                # Use only generate API for completions
                response = requests.post(
                    f"{endpoint}/api/generate",
                    json={
                        "model": model,
                        "prompt": simple_prompt,
                        "stream": False,
                        "options": {
                            "temperature": 0.1,  # Slightly higher for more flexibility
                            "num_predict": 2000,  # Increased for models with thinking mode
                            "num_ctx": int(os.getenv("OLLAMA_CONTEXT_SIZE", "32768")),
                            "stop": ["\n\n\n"],  # Only stop on triple newlines to allow full JSON
                            "top_p": 0.9,
                            "repeat_penalty": 1.0,
                        }
                    },
                    timeout=60
                )
                
                logger.info(f"Ollama response status: {response.status_code}")
                
                if response.status_code == 200:
                    result = response.json()

                    # Check if response was truncated
                    if result.get('done_reason') == 'length':
                        logger.warning(f"⚠️  Ollama response truncated due to token limit for {filename}! Consider increasing num_predict.")

                    # Use only generate API response format
                    ai_response = result.get('response', '').strip()
                    logger.info(f"Using generate API response for {filename}")
                    
                    # Handle thinking field when content is empty - extract JSON if present
                    if not ai_response and 'thinking' in result and result.get('thinking'):
                        thinking_content = result.get('thinking', '')
                        logger.info(f"Content empty, checking thinking field for JSON: '{thinking_content[:200]}...'")
                        
                        # Try to extract JSON from thinking field
                        extracted_json = _extract_json_content_only(thinking_content)
                        if extracted_json:
                            logger.info(f"Found valid JSON in thinking field, using it: {extracted_json}")
                            ai_response = extracted_json
                        else:
                            logger.warning(f"No valid JSON found in thinking field for {filename}")
                    elif 'thinking' in result and result.get('thinking'):
                        logger.info(f"Ollama thinking field (ignored): '{result.get('thinking', '')[:100]}...'")
                    
                    logger.info(f"Raw AI response for {filename}: \"{ai_response[:200]}...\"")
                    logger.info(f"Response length: {len(ai_response)}")
                    
                    # Force clean JSON extraction if the response contains explanatory text
                    if ai_response and ("We need" in ai_response or "Looking at" in ai_response or "The document" in ai_response):
                        logger.warning(f"Response contains explanatory text, attempting JSON extraction")
                        extracted_json = _extract_json_content_only(ai_response)
                        if extracted_json:
                            logger.info(f"Extracted JSON: {extracted_json}")
                            ai_response = extracted_json
                        else:
                            logger.error(f"No valid JSON found in explanatory response")
                            return generate_fallback_suggestions_from_filename(filename, available_controls)
                    
                    if not ai_response:
                        logger.warning(f"AI scanning failed for {filename} - no valid response generated")
                        logger.info(f"Ollama result object: {result}")
                        return []  # Simple empty result instead of filename fallback
                    
                    # Force JSON format validation
                    if not ai_response.strip().startswith('{') or not ai_response.strip().endswith('}'):
                        logger.warning(f"Response doesn't look like JSON for {filename}: {ai_response[:100]}")
                        # Try to extract any JSON from the response
                        ai_response = _extract_json_content_only(ai_response)
                        if not ai_response:
                            logger.error(f"No valid JSON found in response for {filename}")
                            return generate_fallback_suggestions_from_filename(filename, available_controls)
                else:
                    logger.error(f"Ollama error: {response.status_code} - {response.text[:200]}")
                    return generate_fallback_suggestions_from_filename(filename, available_controls)
            else:
                # OpenAI
                model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
                logger.info(f"Using OpenAI model: {model} for document analysis")
                
                try:
                    # First, test with a simple prompt to verify AI client works
                    logger.info(f"Testing AI connectivity with model: {model}")
                    test_response = create_chat_completion_safe(
                        client=ai_client,
                        model=model,
                        messages=[
                            {"role": "user", "content": "Respond with just the JSON: {\"test\": \"success\"}"}
                        ],
                        max_tokens=50,
                        temperature=0.1
                    )
                    test_content = test_response.choices[0].message.content
                    logger.info(f"AI test response: '{test_content}'")
                    logger.info(f"Test response length: {len(test_content) if test_content else 0}")
                    logger.info(f"Response object: {test_response}")
                    
                    if not test_content or not test_content.strip():
                        logger.error(f"AI test failed - empty response! Usage: {test_response.usage}")
                        logger.error(f"Model used: {test_response.model}")
                        return generate_fallback_suggestions_from_filename(filename, available_controls)
                    
                    # Now try the actual analysis
                    response = create_chat_completion_safe(
                        client=ai_client,
                        model=model,
                        messages=[
                            {"role": "system", "content": "You are a compliance expert. Respond only with valid JSON."},
                            {"role": "user", "content": analysis_prompt[:2000]}  # Limit prompt length
                        ],
                        max_tokens=500,  # Reduced tokens
                        temperature=0.3,
                        use_json_mode=True  # Try JSON mode with fallback
                    )
                    ai_response = response.choices[0].message.content
                    logger.info(f"OpenAI analysis response received, length: {len(ai_response) if ai_response else 0}")
                    
                    if not ai_response:
                        logger.warning(f"OpenAI returned empty response for {filename}")
                        logger.info(f"Prompt length was: {len(analysis_prompt)} characters")
                        logger.info(f"Truncated prompt preview: {analysis_prompt[:200]}...")
                        logger.info(f"Available controls: {len(available_controls)}")
                        if available_controls:
                            logger.info(f"Sample control: {available_controls[0]}")
                        # Try with a much simpler approach
                        simple_prompt = f"""You are a JSON API. Analyze '{filename}' and respond ONLY with valid JSON.

Document content: {file_text[:5000] if file_text else 'File analysis'}
Available controls: {[c['code'] for c in available_controls[:5]]}

Respond with ONLY this JSON structure:
{{\"suggestions\": [{{\"control_code\": \"EXACT_CODE\", \"control_title\": \"Full title\", \"framework_name\": \"Framework\", \"confidence\": 0.7, \"reasoning\": \"brief explanation\"}}]}}"""
                        try:
                            simple_response = create_chat_completion_safe(
                                client=ai_client,
                                model=model,
                                messages=[{"role": "user", "content": simple_prompt}],
                                max_tokens=200,
                                temperature=0.3,
                                use_json_mode=True
                            )
                            simple_ai_response = simple_response.choices[0].message.content
                            logger.info(f"Simple prompt response: {simple_ai_response}")
                            if simple_ai_response and simple_ai_response.strip():
                                # Try to parse the simple response
                                try:
                                    simple_parsed = json_module.loads(simple_ai_response)
                                    return simple_parsed.get('suggestions', [])
                                except:
                                    logger.warning("Simple prompt also failed to parse")
                        except Exception as simple_error:
                            logger.error(f"Simple prompt also failed: {simple_error}")
                            return generate_fallback_suggestions_from_filename(filename, available_controls)
                    
                except Exception as openai_error:
                    logger.error(f"OpenAI API error for {filename}: {type(openai_error).__name__}: {openai_error}")
                    
                    # Check for specific error types
                    error_str = str(openai_error).lower()
                    if 'api key' in error_str or 'authentication' in error_str:
                        logger.error("❌ API Key Issue: Check your OpenAI API key in settings")
                    elif 'quota' in error_str or 'billing' in error_str:
                        logger.error("💰 Quota Issue: Check your OpenAI billing/credits")
                    elif 'rate limit' in error_str:
                        logger.error("🚦 Rate Limited: Too many requests to OpenAI")
                    elif 'model' in error_str:
                        logger.error(f"🤖 Model Issue: Model '{model}' might not be available")
                    else:
                        logger.error(f"🔧 Unknown OpenAI Error: {openai_error}")
                    
                    return generate_fallback_suggestions_from_filename(filename, available_controls)
            
            # Parse AI response with improved error handling
            try:
                # Log the raw response for debugging
                logger.info(f"Raw AI response for {filename}: {repr(ai_response[:500])}")
                
                # Check if response looks like an error message
                if ai_response and ("error" in ai_response.lower() or "internal server" in ai_response.lower() or ai_response.startswith("HTTP/")):
                    logger.warning(f"AI response appears to be an error message: {ai_response[:200]}")
                    return generate_fallback_suggestions_from_filename(filename, available_controls)
                
                # Clean the AI response - remove any markdown formatting or extra text
                cleaned_response = ai_response.strip() if ai_response else ""
                
                if not cleaned_response:
                    logger.warning(f"Empty AI response for {filename}")
                    return generate_fallback_suggestions_from_filename(filename, available_controls)
                
                # Look for JSON content within the response
                if '```json' in cleaned_response:
                    # Extract JSON from markdown code blocks
                    start = cleaned_response.find('```json') + 7
                    end = cleaned_response.find('```', start)
                    if end != -1:
                        cleaned_response = cleaned_response[start:end].strip()
                elif '```' in cleaned_response:
                    # Extract from general code blocks
                    start = cleaned_response.find('```') + 3
                    end = cleaned_response.rfind('```')
                    if end != -1 and end > start:
                        cleaned_response = cleaned_response[start:end].strip()
                
                # Find JSON object in the response
                json_start = cleaned_response.find('{')
                json_end = cleaned_response.rfind('}') + 1
                
                if json_start != -1 and json_end > json_start:
                    json_content = cleaned_response[json_start:json_end]
                    try:
                        parsed_response = json_module.loads(json_content)
                    except json_module.JSONDecodeError as json_err:
                        logger.warning(f"JSON parsing failed for extracted content: {json_err}")
                        logger.info(f"Extracted JSON content: {repr(json_content[:300])}")
                        return generate_fallback_suggestions_from_filename(filename, available_controls)
                else:
                    try:
                        parsed_response = json_module.loads(cleaned_response)
                    except json_module.JSONDecodeError as json_err:
                        logger.warning(f"JSON parsing failed for full response: {json_err}")
                        logger.info(f"Cleaned response: {repr(cleaned_response[:300])}")
                        return generate_fallback_suggestions_from_filename(filename, available_controls)
                    
                suggestions = parsed_response.get('suggestions', [])
                
                # Validate and return suggestions
                valid_suggestions = []
                for suggestion in suggestions[:3]:
                    if isinstance(suggestion, dict) and all(key in suggestion for key in ['control_code', 'control_title']):
                        confidence = suggestion.get('confidence', 0.5)
                        # Handle various confidence formats
                        if isinstance(confidence, str):
                            try:
                                confidence = float(confidence.replace('%', '')) / 100 if '%' in confidence else float(confidence)
                            except:
                                confidence = 0.5
                        
                        valid_suggestions.append({
                            'control_code': str(suggestion['control_code']),
                            'control_title': str(suggestion['control_title']),
                            'framework_name': str(suggestion.get('framework_name', 'Unknown')),
                            'confidence': max(0.0, min(1.0, float(confidence))),
                            'reasoning': str(suggestion.get('reasoning', 'AI analysis'))
                        })
                
                return valid_suggestions
                
            except (json_module.JSONDecodeError, ValueError, KeyError) as e:
                logger.warning(f"Failed to parse AI response for {filename}: {str(e)[:100]}")
                logger.info(f"AI response content (first 300 chars): {repr(ai_response[:300])}")
                
                # Check if response is empty or whitespace
                if not ai_response or not ai_response.strip():
                    logger.warning(f"AI returned empty response for {filename}")
                    return generate_fallback_suggestions_from_filename(filename, available_controls)
                
                # Fallback: try to extract control suggestions from free text
                fallback_suggestions = extract_suggestions_from_text(ai_response, available_controls)
                if fallback_suggestions:
                    logger.info(f"Extracted {len(fallback_suggestions)} suggestions from text for {filename}")
                    return fallback_suggestions
                
                # Final fallback to filename analysis
                return generate_fallback_suggestions_from_filename(filename, available_controls)
                
        except Exception as e:
            logger.error(f"AI analysis completely failed for {filename}: {e}")
            # Return smart fallback suggestions based on filename
            return generate_fallback_suggestions_from_filename(filename, available_controls)
            
    except Exception as e:
        logger.error(f"Content analysis error for {file.filename}: {e}")
        logger.exception("Full content analysis error:")
        return []

# Ensure the function always returns a list
async def safe_analyze_file_content_for_controls(file, file_content):
    """Wrapper for analyze_file_content_for_controls that ensures it always returns a list."""
    try:
        result = await analyze_file_content_for_controls(file, file_content)
        if not isinstance(result, list):
            logger.warning(f"analyze_file_content_for_controls returned non-list: {type(result)}")
            return []
        
        # Enforce one-control-per-document policy
        if len(result) > 1:
            logger.info(f"Multiple controls suggested for {file.filename}, returning only the highest confidence match")
            # Sort by confidence and return only the highest one
            result_sorted = sorted(result, key=lambda x: x.get('confidence', 0), reverse=True)
            return [result_sorted[0]]
        
        return result
    except Exception as e:
        logger.error(f"Safe analysis wrapper caught error for {file.filename}: {e}")
        return []

def find_unprocessed_documents() -> List[Document]:
    """Find documents that haven't been AI processed yet or were created more than 1 hour ago without control links."""
    try:
        db = next(get_db())
        
        # Find documents that either:
        # 1. Have no control links at all (never processed)
        # 2. Were created more than 1 hour ago and still have no control links (likely failed)
        one_hour_ago = datetime.utcnow() - timedelta(hours=1)
        
        # Find documents without any control links OR created over an hour ago
        unprocessed = db.query(Document).outerjoin(DocumentControlLink).filter(
            Document.created_at < one_hour_ago  # Created over 1 hour ago
        ).group_by(Document.id).having(
            db.func.count(DocumentControlLink.id) == 0  # No associated control links
        ).all()
        
        db.close()
        logger.info(f"Found {len(unprocessed)} unprocessed documents")
        return unprocessed
        
    except Exception as e:
        logger.error(f"Error finding unprocessed documents: {e}")
        return []

async def retry_ai_processing():
    """Retry AI processing for unprocessed documents."""
    try:
        logger.info("Starting hourly AI processing retry check...")
        unprocessed_docs = find_unprocessed_documents()
        
        for doc in unprocessed_docs:
            try:
                logger.info(f"Retrying AI processing for document: {doc.id} - {doc.filename}")
                
                # Download the document content from storage
                file_content = storage.download(doc.storage_key)
                
                # Process in background
                await process_document_ai_analysis_background(str(doc.id), doc.filename, file_content)
                
            except Exception as e:
                logger.error(f"Failed to retry AI processing for document {doc.id}: {e}")
                
    except Exception as e:
        logger.error(f"Error during AI processing retry: {e}")

async def periodic_ai_retry_task():
    """Background task that runs every hour to retry AI processing."""
    while True:
        try:
            await asyncio.sleep(3600)  # Wait 1 hour
            await retry_ai_processing()
        except Exception as e:
            logger.error(f"Error in periodic AI retry task: {e}")
            await asyncio.sleep(300)  # Wait 5 minutes before trying again

async def process_document_ai_analysis_background(document_id: str, filename: str, file_content: bytes):
    """Process AI analysis in background and store results."""
    try:
        logger.info(f"Starting background AI analysis for document {document_id}: {filename}")
        
        # Create a mock file object for the analysis
        class MockFile:
            def __init__(self, filename, content):
                self.filename = filename
                self._content = content
                self._position = 0
                # Determine content type from filename
                if filename.lower().endswith('.txt'):
                    self.content_type = 'text/plain'
                elif filename.lower().endswith('.pdf'):
                    self.content_type = 'application/pdf'
                elif filename.lower().endswith('.docx'):
                    self.content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                elif filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                    self.content_type = 'image/' + filename.lower().split('.')[-1]
                else:
                    self.content_type = 'application/octet-stream'
            
            async def seek(self, position):
                self._position = position
            
            async def read(self):
                return self._content
        
        mock_file = MockFile(filename, file_content)
        
        # Perform the AI analysis
        suggested_controls = await safe_analyze_file_content_for_controls(mock_file, file_content)
        
        if not suggested_controls:
            # Use filename fallback if AI analysis fails - get available controls from database
            try:
                from sqlalchemy.orm import sessionmaker
                from database import engine
                from models import Control
                SessionLocal = sessionmaker(bind=engine)
                db = SessionLocal()
                
                controls = db.query(Control).all()
                available_controls = [
                    {
                        'code': control.code,
                        'title': control.title,
                        'framework': getattr(control, 'framework_name', 'Essential Eight')
                    }
                    for control in controls
                ]
                db.close()
                
                suggested_controls = generate_fallback_suggestions_from_filename(filename, available_controls)[:1]
                logger.info(f"Using filename fallback for {filename}: {len(suggested_controls)} suggestions")
            except Exception as e:
                logger.error(f"Failed to get fallback suggestions for {filename}: {e}")
                suggested_controls = []
        
        logger.info(f"Background AI analysis completed for {filename}: {len(suggested_controls)} suggestions")
        
        # Store results in database for later retrieval
        try:
            from sqlalchemy.orm import sessionmaker
            from database import engine
            from models import Document, DocumentControlLink, Control
            SessionLocal = sessionmaker(bind=engine)
            db = SessionLocal()
            
            # Get confidence threshold and dual vision settings
            from models import Settings
            settings = db.query(Settings).filter(Settings.id == 1).first()
            min_threshold = settings.min_confidence_threshold if settings else 0.90
            use_dual_vision = bool(settings.use_dual_vision_validation) if settings and hasattr(settings, 'use_dual_vision_validation') else False

            # Log which mode we're using
            if use_dual_vision:
                logger.info(f"📸 DUAL VISION MODE: Will validate with both GPT-4o AND Qwen2-VL")
            else:
                logger.info(f"📸 SINGLE MODEL MODE: Using configured provider ({settings.ai_provider if settings else 'default'})")

            if use_dual_vision and suggested_controls:
                logger.info(f"Dual vision validation enabled for document {document_id}")
                try:
                    # Re-analyze with both vision models
                    from ai_scanner import get_vision_clients_for_dual_validation
                    vision_clients = get_vision_clients_for_dual_validation()

                    if len(vision_clients) >= 2:
                        logger.info(f"Running dual vision validation with {len(vision_clients)} models")

                        # Re-run analysis with both models
                        dual_results = []
                        for provider_name, client_info in vision_clients.items():
                            try:
                                logger.info(f"Analyzing with {provider_name} - {client_info['model']}")
                                # Re-analyze the document with this specific client
                                mock_file = MockFile(filename, file_content)
                                result = await safe_analyze_file_content_for_controls(mock_file, file_content)
                                if result:
                                    dual_results.append({
                                        'provider': provider_name,
                                        'model': client_info['model'],
                                        'result': result[0] if result else None
                                    })
                            except Exception as e:
                                logger.error(f"Dual vision analysis failed for {provider_name}: {e}")

                        # Compare results - only proceed if both models agree
                        if len(dual_results) == 2:
                            result1 = dual_results[0]['result']
                            result2 = dual_results[1]['result']

                            if result1 and result2:
                                control1 = result1.get('control_code')
                                control2 = result2.get('control_code')
                                conf1 = result1.get('confidence', 0.0)
                                conf2 = result2.get('confidence', 0.0)

                                if control1 == control2:
                                    # Both models agree! Use minimum confidence (most conservative)
                                    min_conf = min(conf1, conf2)
                                    logger.info(f"✓ Dual vision CONSENSUS: Both models agree on {control1} (conf: {conf1:.2f} & {conf2:.2f}, using min: {min_conf:.2f})")

                                    # Update suggested_controls with validated result
                                    suggested_controls = [{
                                        'control_code': control1,
                                        'confidence': min_conf,
                                        'reasoning': f"Dual validation: {dual_results[0]['model']} ({conf1:.2f}) and {dual_results[1]['model']} ({conf2:.2f}) both agree"
                                    }]
                                else:
                                    # Models disagree - reject the link for safety
                                    logger.warning(f"✗ Dual vision DISAGREEMENT: {dual_results[0]['model']} suggested {control1} ({conf1:.2f}), {dual_results[1]['model']} suggested {control2} ({conf2:.2f}) - NO LINK CREATED")
                                    suggested_controls = []  # Clear suggestions - no consensus
                    else:
                        logger.warning(f"Dual vision validation requires 2 models, only {len(vision_clients)} available - falling back to single model")
                except Exception as e:
                    logger.error(f"Dual vision validation error: {e}")
                    # Fall back to single model result on error
            else:
                # Single model mode - use the initial analysis result
                if suggested_controls:
                    logger.info(f"Using single model result: {suggested_controls[0].get('control_code')} (confidence: {suggested_controls[0].get('confidence', 0.0):.2f})")

            # Update document with AI processing complete status
            document = db.query(Document).filter(Document.id == document_id).first()
            if document and suggested_controls:
                # Find the suggested control in database
                control = db.query(Control).filter(Control.code == suggested_controls[0]['control_code']).first()
                if control:
                    confidence = suggested_controls[0].get('confidence', 0.0)

                    # Only create link if confidence meets threshold
                    if confidence >= min_threshold:
                        # Create document-control link
                        existing_link = db.query(DocumentControlLink).filter_by(
                            document_id=document_id,
                            control_id=control.id
                        ).first()

                        if not existing_link:
                            link = DocumentControlLink(
                                document_id=document_id,
                                control_id=control.id,
                                confidence=confidence,
                                reasoning=suggested_controls[0].get('reasoning', 'AI analysis')
                            )
                            db.add(link)
                            db.commit()
                            logger.info(f"Created document-control link: {document.filename} → {control.code} (confidence: {confidence:.2f})")
                    else:
                        logger.info(f"Skipped link for {document.filename} → {control.code}: confidence {confidence:.2f} below threshold {min_threshold}")
            
            db.close()
        except Exception as e:
            logger.error(f"Failed to store AI results for document {document_id}: {e}")
        
        logger.info(f"AI suggestions for document {document_id}: {suggested_controls}")
        
        return suggested_controls
        
    except Exception as e:
        logger.error(f"Background AI analysis failed for document {document_id}: {e}")
        logger.exception("Background AI analysis error details:")
        return []

def extract_suggestions_from_text(text_response: str, available_controls: list) -> list:
    """Extract control suggestions from free-form AI text when JSON parsing fails."""
    suggestions = []
    text_lower = text_response.lower()
    
    # Look for mentioned control codes in the response
    for control in available_controls[:50]:  # Check first 50 controls with larger context
        control_code = control['code'].lower()
        control_title = control['title'].lower()
        
        # Check if control is mentioned in the response
        if (control_code in text_lower or 
            any(word in text_lower for word in control_title.split() if len(word) > 3)):
            
            # Estimate confidence based on how prominently it's mentioned (more strict)
            confidence = 0.4  # Lower base confidence
            if control_code in text_lower:
                confidence = 0.6  # Reduced from 0.8
            if 'relevant' in text_lower or 'applicable' in text_lower:
                confidence += 0.1
            if 'not' in text_lower and control_code in text_lower:
                confidence = max(0.2, confidence - 0.4)  # More penalty for negation
                
            suggestions.append({
                'control_code': control['code'],
                'control_title': control['title'],
                'framework_name': control.get('framework', 'Unknown'),
                'confidence': min(0.9, confidence),
                'reasoning': f'AI mentioned this control in analysis text'
            })
            
            if len(suggestions) >= 2:  # Limit fallback suggestions
                break
    
    return suggestions

def generate_fallback_suggestions_from_filename(filename: str, available_controls: list) -> list:
    """Generate basic suggestions based on filename when AI analysis completely fails."""
    suggestions = []
    filename_lower = filename.lower()
    
    # Specific filename patterns for Essential Eight controls
    specific_patterns = {
        # EE-8 Regular Backups
        ('backup', 'recovery', 'restore', '08_backup'): ('EE-8', 'Regular Backups', 0.9),
        # EE-7 Multi-Factor Authentication  
        ('multi_factor', 'mfa', '2fa', '07_multi', 'authentication'): ('EE-7', 'Multi-Factor Authentication', 0.9),
        # EE-5 Administrative Privileges
        ('05_administrative', 'privilege', 'admin'): ('EE-5', 'Restrict Administrative Privileges', 0.9),
        # EE-1 Application Control
        ('01_application', 'app_control', 'software_control'): ('EE-1', 'Application Control', 0.9),
        # EE-2 Patch Applications
        ('02_patch', 'app_patch', 'application_patch'): ('EE-2', 'Patch Applications', 0.9),
        # EE-6 Patch Operating Systems
        ('06_patch', 'os_patch', 'system_patch', 'update', 'windows_update', 'os_update'): ('EE-6', 'Patch Operating Systems', 0.9),
        # EE-3 Configure Microsoft Office Macro Settings
        ('03_macro', 'office_macro', 'macro_settings', 'macro', 'macros', 'marco', 'vba', 'trust_center', 'office_security'): ('EE-3', 'Configure Microsoft Office Macro Settings', 0.9),
        # EE-4 User Application Hardening
        ('04_hardening', 'browser', 'user_app'): ('EE-4', 'User Application Hardening', 0.9),
    }
    
    # Check specific patterns first (highest priority)
    for keywords, (code, title, confidence) in specific_patterns.items():
        if any(keyword in filename_lower for keyword in keywords):
            # Find the matching control in available_controls
            for control in available_controls:
                if control.get('code') == code or code in control.get('title', ''):
                    suggestions.append({
                        'control_code': control['code'],
                        'control_title': control['title'],
                        'framework_name': control.get('framework', 'Essential Eight'),
                        'confidence': confidence,
                        'reasoning': f'Filename indicates this is a {title.lower()} policy document'
                    })
                    return suggestions  # Return immediately for specific matches
    
    # Fallback to generic keyword mapping  
    generic_keywords_mapping = {
        ('mfa', 'multi-factor', '2fa', 'authentication', 'auth'): 'authentication',
        ('access', 'identity', 'user', 'login'): 'access',
        ('policy', 'procedure', 'governance'): 'policy', 
        ('security', 'incident', 'response', 'error'): 'security',
        ('config', 'configuration', 'setting'): 'configuration',
        ('log', 'audit', 'monitoring', 'compliance'): 'audit'
    }
    
    for keywords, category in generic_keywords_mapping.items():
        if any(keyword in filename_lower for keyword in keywords):
            # Find matching controls
            for control in available_controls:
                control_lower = (control['title'] + ' ' + control.get('description', '')).lower()
                if (category in control_lower or
                    any(keyword in control_lower for keyword in keywords[:2])):

                    confidence = 0.4 if category == 'authentication' and 'mfa' in filename_lower else 0.35  # Lower baseline for filename matching
                    
                    suggestions.append({
                        'control_code': control['code'],
                        'control_title': control['title'],
                        'framework_name': control.get('framework', 'Unknown'),
                        'confidence': confidence,
                        'reasoning': f'Filename suggests {category}-related content'
                    })
                    
                    if len(suggestions) >= 2:
                        return suggestions
    
    return suggestions

@app.get("/")
async def root():
    return {"message": "GeekyGoose Compliance API is running"}

@app.get("/health")
async def health():
    return {"status": "healthy"}

@app.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    # Validate file type
    if file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"File type {file.content_type} not allowed. Allowed types: {', '.join(ALLOWED_MIME_TYPES)}"
        )
    
    # Check file size
    file_content = await file.read()
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size: {MAX_FILE_SIZE // (1024*1024)}MB"
        )
    
    # Reset file position
    await file.seek(0)
    
    try:
        # Upload to storage
        storage_key, sha256_hash, file_size = storage.upload_file(
            file.file, file.filename, file.content_type
        )
        
        # For demo purposes, create a default org and user if they don't exist
        org = db.query(Org).first()
        if not org:
            org = Org(id=uuid.uuid4(), name="Demo Organization")
            db.add(org)
            db.commit()
            db.refresh(org)
        
        user = db.query(User).filter(User.org_id == org.id).first()
        if not user:
            user = User(
                id=uuid.uuid4(),
                org_id=org.id,
                email="demo@example.com",
                name="Demo User"
            )
            db.add(user)
            db.commit()
            db.refresh(user)
        
        # Save to database
        document = Document(
            org_id=org.id,
            filename=file.filename,
            mime_type=file.content_type,
            storage_key=storage_key,
            file_size=file_size,
            uploaded_by=user.id,
            sha256=sha256_hash
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)

        # Trigger text extraction task (required for compliance scanning)
        try:
            from worker_tasks import extract_document_text
            extract_task = extract_document_text.delay(str(document.id))
            logger.info(f"Triggered text extraction task for {file.filename}: {extract_task.id}")
        except Exception as e:
            logger.error(f"Failed to trigger text extraction for {file.filename}: {e}")

        # Return immediate response without AI analysis to prevent timeouts
        # AI analysis will be processed in background
        suggested_controls = []

        # Start background AI analysis (but don't wait for it)
        import asyncio
        try:
            # Schedule AI analysis in background
            asyncio.create_task(process_document_ai_analysis_background(
                document.id,
                file.filename,
                file_content
            ))
            logger.info(f"Scheduled background AI analysis for {file.filename}")
        except Exception as e:
            logger.error(f"Failed to schedule background AI analysis for {file.filename}: {e}")
        
        # Provide immediate filename-based suggestion for quick feedback
        try:
            # Get available controls for immediate suggestions
            controls = db.query(Control).all()
            available_controls = [
                {
                    'code': control.code,
                    'title': control.title,
                    'framework': getattr(control, 'framework_name', 'Essential Eight')
                }
                for control in controls
            ]
            
            suggested_controls = generate_fallback_suggestions_from_filename(file.filename, available_controls)[:1]
            logger.info(f"Providing immediate filename-based suggestions for {file.filename}: {len(suggested_controls)} suggestions")
        except Exception as e:
            logger.error(f"Filename-based suggestions failed for {file.filename}: {e}")
            suggested_controls = []
        
        return {
            "id": str(document.id),
            "filename": document.filename,
            "mime_type": document.mime_type,
            "file_size": document.file_size,
            "sha256": document.sha256,
            "created_at": document.created_at.isoformat(),
            "download_url": storage.get_download_url(storage_key),
            "suggested_controls": suggested_controls
        }
        
    except Exception as e:
        logger.error(f"Upload failed for {file.filename if file else 'unknown file'}: {e}")
        logger.exception("Full upload error details:")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.get("/documents")
async def list_documents(db: Session = Depends(get_db)):
    documents = db.query(Document).order_by(Document.created_at.desc()).all()
    
    result = []
    for doc in documents:
        # Check if document has control links (AI processing complete)
        links = db.query(DocumentControlLink).filter(DocumentControlLink.document_id == doc.id).all()
        
        result.append({
            "id": str(doc.id),
            "filename": doc.filename,
            "mime_type": doc.mime_type,
            "file_size": doc.file_size,
            "sha256": doc.sha256,
            "created_at": doc.created_at.isoformat(),
            "download_url": f"/api/documents/{doc.id}/download",  # Fixed download URL
            "ai_processed": len(links) > 0,
            "control_links": [
                {
                    "control_id": str(link.control_id),
                    "control_code": link.control.code if link.control else "Unknown",
                    "control_title": link.control.title if link.control else "Unknown",
                    "confidence": link.confidence,
                    "reasoning": link.reasoning
                }
                for link in links
            ]
        })
    
    return {"documents": result}

@app.get("/documents/{document_id}/download")
async def download_document(document_id: str, db: Session = Depends(get_db)):
    """Download a document file."""
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        logger.info(f"Attempting to download document {document_id}: {document.filename}")
        logger.info(f"Storage key: {document.storage_key}")
        logger.info(f"MIME type: {document.mime_type}")
        
        # Get file content from storage
        file_content = storage.download_file(document.storage_key)
        
        if not file_content:
            logger.error(f"No file content retrieved for document {document_id}")
            raise HTTPException(status_code=404, detail="File content not found")
        
        logger.info(f"Retrieved file content, size: {len(file_content)} bytes")
        
        from fastapi.responses import Response
        return Response(
            content=file_content,
            media_type=document.mime_type,
            headers={
                "Content-Disposition": f"attachment; filename=\"{document.filename}\""
            }
        )
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        logger.error(f"Download failed for document {document_id}: {type(e).__name__}: {e}")
        logger.error(f"Document details: filename={document.filename}, storage_key={document.storage_key}")
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

@app.post("/reports/comprehensive-analysis")
async def run_comprehensive_ai_analysis(request: dict, db: Session = Depends(get_db)):
    """Run comprehensive AI analysis across all documents and controls."""
    try:
        framework_id = request.get('framework_id')
        if not framework_id:
            raise HTTPException(status_code=400, detail="Framework ID required")
        
        logger.info(f"Starting comprehensive AI analysis for framework {framework_id}")
        
        # Get all documents (AI processing status determined by control links)
        documents = db.query(Document).all()
        
        # Get all controls for the framework
        controls = db.query(Control).filter(Control.framework_id == framework_id).all()
        
        # Get all document-control links
        control_links = db.query(DocumentControlLink).all()
        
        # Filter documents that have AI processing (have control links)
        ai_processed_document_ids = set(link.document_id for link in control_links)
        ai_processed_documents = [doc for doc in documents if doc.id in ai_processed_document_ids]
        
        # Calculate comprehensive metrics
        total_controls = len(controls)
        controls_with_evidence = len(set(link.control_id for link in control_links))
        coverage_percentage = round((controls_with_evidence / total_controls * 100) if total_controls > 0 else 0)
        
        # Calculate average confidence
        if control_links:
            avg_confidence = round(sum(link.confidence for link in control_links) / len(control_links) * 100)
        else:
            avg_confidence = 0
        
        # Identify high-risk gaps (controls with no evidence or low confidence)
        high_risk_gaps = 0
        control_analysis = []
        
        for control in controls:
            control_evidence = [link for link in control_links if link.control_id == control.id]
            
            if not control_evidence:
                high_risk_gaps += 1
                control_analysis.append({
                    'control_code': control.code,
                    'control_title': control.title,
                    'risk_level': 'HIGH',
                    'issue': 'No evidence found',
                    'evidence_count': 0,
                    'avg_confidence': 0
                })
            else:
                max_confidence = max(link.confidence for link in control_evidence)
                if max_confidence < 0.7:  # Raised threshold - require stronger evidence to not be high risk
                    high_risk_gaps += 1
                    control_analysis.append({
                        'control_code': control.code,
                        'control_title': control.title,
                        'risk_level': 'HIGH',
                        'issue': 'Low confidence evidence',
                        'evidence_count': len(control_evidence),
                        'avg_confidence': round(max_confidence * 100)
                    })
        
        # Generate AI recommendations based on analysis
        recommendations = []
        
        if coverage_percentage < 50:
            recommendations.append("Coverage is below 50%. Consider uploading more compliance documentation.")
        
        if avg_confidence < 70:
            recommendations.append("Average confidence is low. Review evidence quality and ensure documents contain specific compliance details.")
        
        if high_risk_gaps > total_controls * 0.3:
            recommendations.append("High number of gaps detected. Prioritize evidence collection for missing controls.")
        
        # Document type analysis
        document_types = {}
        for doc in documents:
            doc_type = doc.mime_type
            if doc_type not in document_types:
                document_types[doc_type] = {'count': 0, 'with_links': 0}
            document_types[doc_type]['count'] += 1
            if doc.id in ai_processed_document_ids:
                document_types[doc_type]['with_links'] += 1
        
        # Evidence gaps by control
        missing_controls = [c.code for c in controls if c.id not in [link.control_id for link in control_links]]
        
        if missing_controls:
            recommendations.append(f"Missing evidence for controls: {', '.join(missing_controls[:5])}{'...' if len(missing_controls) > 5 else ''}")
        
        if len(ai_processed_documents) < 5:
            recommendations.append("Consider uploading more supporting documents for comprehensive compliance coverage.")
        
        # Run AI analysis on the overall compliance state
        from ai_scanner import get_ai_client
        ai_client = get_ai_client()
        if ai_client and not isinstance(ai_client, dict):
            try:
                analysis_prompt = f"""
                Analyze this compliance state and provide strategic recommendations:
                
                Framework: {controls[0].framework.name if controls else 'Unknown'}
                Total Controls: {total_controls}
                Coverage: {coverage_percentage}%
                Documents Analyzed: {len(ai_processed_documents)}
                Control Links: {len(control_links)}
                High Risk Gaps: {high_risk_gaps}
                
                Missing Controls: {', '.join(missing_controls[:10])}
                
                Provide 3-5 specific, actionable recommendations for improving compliance posture.
                Focus on prioritization and practical next steps.
                """
                
                ai_response = create_chat_completion_safe(
                    client=ai_client,
                    model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
                    messages=[
                        {"role": "system", "content": "You are a compliance expert. Provide concise, actionable recommendations."},
                        {"role": "user", "content": analysis_prompt}
                    ],
                    max_tokens=500,
                    temperature=0.1
                )
                
                if ai_response and ai_response.choices[0].message.content:
                    ai_recommendations = ai_response.choices[0].message.content.strip().split('\n')
                    # Clean and filter AI recommendations
                    ai_recommendations = [rec.strip('- •').strip() for rec in ai_recommendations if rec.strip() and len(rec.strip()) > 10]
                    recommendations.extend(ai_recommendations[:5])
                    
            except Exception as e:
                logger.error(f"AI recommendation generation failed: {e}")
        
        # Ensure we have at least some recommendations
        if not recommendations:
            recommendations = [
                "Continue monitoring compliance status and uploading relevant documentation.",
                "Review control requirements and ensure adequate evidence is available.",
                "Consider conducting periodic compliance assessments."
            ]
        
        result = {
            'framework_id': framework_id,
            'analysis_timestamp': datetime.utcnow().isoformat(),
            'coverage_percentage': coverage_percentage,
            'total_controls': total_controls,
            'controls_with_evidence': controls_with_evidence,
            'total_documents': len(ai_processed_documents),
            'total_control_links': len(control_links),
            'avg_confidence': avg_confidence,
            'high_risk_gaps': high_risk_gaps,
            'recommendations': recommendations[:8],  # Limit to 8 recommendations
            'document_types': document_types,
            'control_analysis': control_analysis[:10],  # Top 10 high-risk controls
            'missing_controls': missing_controls[:20]  # Top 20 missing controls
        }
        
        logger.info(f"Comprehensive analysis complete: {coverage_percentage}% coverage, {high_risk_gaps} high-risk gaps")
        return result
        
    except Exception as e:
        logger.error(f"Comprehensive analysis failed: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@app.get("/documents/{document_id}/ai-status")
async def get_document_ai_status(document_id: str, db: Session = Depends(get_db)):
    """Check if AI processing is complete for a document."""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            # Document was deleted, return completed status to stop polling
            return {
                "document_id": document_id,
                "filename": "Deleted Document",
                "ai_processed": True,  # Mark as processed to stop polling
                "control_links": [],
                "deleted": True
            }
        
        # Check if document has control links (AI processing complete)
        links = db.query(DocumentControlLink).filter(DocumentControlLink.document_id == document_id).all()
        
        return {
            "document_id": document_id,
            "filename": document.filename,
            "ai_processed": len(links) > 0,
            "control_links": [
                {
                    "control_id": str(link.control_id),
                    "control_code": link.control.code if link.control else "Unknown", 
                    "control_title": link.control.title if link.control else "Unknown",
                    "confidence": link.confidence,
                    "reasoning": link.reasoning
                }
                for link in links
            ]
        }
    except Exception as e:
        logger.error(f"Failed to get AI status for document {document_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to get AI status")

@app.post("/admin/retry-ai-processing")
async def manual_retry_ai_processing():
    """Manually trigger AI processing retry for unprocessed documents."""
    try:
        await retry_ai_processing()
        return {"status": "success", "message": "AI processing retry triggered"}
    except Exception as e:
        logger.error(f"Failed to trigger AI processing retry: {e}")
        raise HTTPException(status_code=500, detail="Failed to trigger AI processing retry")

@app.delete("/documents/{document_id}")
async def delete_document(document_id: str, db: Session = Depends(get_db)):
    document = db.query(Document).filter(Document.id == document_id).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # Delete from storage first
        storage.delete_file(document.storage_key)
        
        # Delete related records first to avoid foreign key constraint violations
        # Delete document control links
        db.query(DocumentControlLink).filter(DocumentControlLink.document_id == document_id).delete()
        
        # Delete evidence links
        db.query(EvidenceLink).filter(EvidenceLink.document_id == document_id).delete()
        
        # Delete document pages (handled by relationship cascade, but explicit is better)
        db.query(DocumentPage).filter(DocumentPage.document_id == document_id).delete()
        
        # Finally delete the document
        db.delete(document)
        db.commit()
        
        return {"message": "Document deleted successfully"}
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Delete failed: {str(e)}")

# Pydantic models for API requests
class LinkEvidenceRequest(BaseModel):
    control_id: str
    requirement_id: Optional[str] = None
    note: Optional[str] = None

class CreateScanRequest(BaseModel):
    control_id: str

# Frameworks and Controls endpoints
@app.get("/frameworks")
async def list_frameworks(db: Session = Depends(get_db)):
    """List all available compliance frameworks."""
    frameworks = db.query(Framework).all()
    return {
        "frameworks": [
            {
                "id": str(f.id),
                "name": f.name,
                "version": f.version,
                "description": f.description,
                "created_at": f.created_at.isoformat()
            }
            for f in frameworks
        ]
    }

@app.get("/frameworks/{framework_id}/controls")
async def list_controls(framework_id: str, db: Session = Depends(get_db)):
    """List all controls for a framework."""
    controls = db.query(Control).filter(Control.framework_id == framework_id).all()
    
    result = []
    for control in controls:
        requirements = db.query(Requirement).filter(Requirement.control_id == control.id).all()
        # Count linked documents for this control
        linked_docs_count = db.query(DocumentControlLink).filter(DocumentControlLink.control_id == control.id).count()
        
        result.append({
            "id": str(control.id),
            "code": control.code,
            "title": control.title,
            "description": control.description,
            "requirements_count": len(requirements),
            "linked_documents_count": linked_docs_count,
            "created_at": control.created_at.isoformat()
        })
    
    return {"controls": result}

def get_control_by_id_or_code(db: Session, control_identifier: str) -> Control:
    """
    Find a control by either UUID or code (case-insensitive).
    Returns the control or None if not found.
    """
    from uuid import UUID

    # Try to parse as UUID first
    try:
        uuid_obj = UUID(control_identifier)
        control = db.query(Control).filter(Control.id == uuid_obj).first()
        if control:
            return control
    except (ValueError, AttributeError):
        pass

    # Try as code (case-insensitive)
    control = db.query(Control).filter(Control.code.ilike(control_identifier)).first()
    return control

@app.get("/controls/{control_id}")
async def get_control_details(control_id: str, db: Session = Depends(get_db)):
    """Get detailed information about a specific control."""
    control = get_control_by_id_or_code(db, control_id)
    if not control:
        raise HTTPException(status_code=404, detail="Control not found")
    
    requirements = db.query(Requirement).filter(Requirement.control_id == control.id).all()
    
    return {
        "id": str(control.id),
        "framework_id": str(control.framework_id),
        "framework_name": control.framework.name,
        "code": control.code,
        "title": control.title,
        "description": control.description,
        "requirements": [
            {
                "id": str(req.id),
                "req_code": req.req_code,
                "text": req.text,
                "maturity_level": req.maturity_level,
                "guidance": req.guidance
            }
            for req in requirements
        ],
        "created_at": control.created_at.isoformat()
    }

@app.post("/documents/{document_id}/link-evidence")
async def link_evidence_to_control(
    document_id: str,
    request: LinkEvidenceRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    """Link a document as evidence for a control/requirement."""
    
    # Verify document exists
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Verify control exists
    control = db.query(Control).filter(Control.id == request.control_id).first()
    if not control:
        raise HTTPException(status_code=404, detail="Control not found")
    
    # Verify requirement exists if provided
    requirement = None
    if request.requirement_id:
        requirement = db.query(Requirement).filter(Requirement.id == request.requirement_id).first()
        if not requirement:
            raise HTTPException(status_code=404, detail="Requirement not found")
    
    # Check if link already exists
    existing = db.query(EvidenceLink).filter(
        EvidenceLink.document_id == document_id,
        EvidenceLink.control_id == request.control_id,
        EvidenceLink.requirement_id == request.requirement_id
    ).first()
    
    if existing:
        raise HTTPException(status_code=400, detail="Evidence link already exists")
    
    # Create evidence link
    evidence_link = EvidenceLink(
        org_id=document.org_id,
        control_id=request.control_id,
        requirement_id=request.requirement_id,
        document_id=document_id,
        note=request.note
    )
    
    db.add(evidence_link)
    db.commit()
    
    # Trigger text extraction in background if not already done
    background_tasks.add_task(extract_document_text.delay, str(document.id))
    
    return {
        "message": "Evidence linked successfully",
        "link_id": str(evidence_link.id)
    }

@app.get("/controls/{control_id}/evidence")
async def get_control_evidence(control_id: str, db: Session = Depends(get_db)):
    """Get all evidence linked to a control (both AI-linked and manual)."""

    # Find the control by ID or code
    control = get_control_by_id_or_code(db, control_id)
    if not control:
        raise HTTPException(status_code=404, detail="Control not found")

    result = []

    # Get AI-linked evidence (DocumentControlLink)
    ai_links = db.query(DocumentControlLink).filter(
        DocumentControlLink.control_id == control.id
    ).all()

    for link in ai_links:
        result.append({
            "id": str(link.id),
            "document": {
                "id": str(link.document.id),
                "filename": link.document.filename,
                "mime_type": link.document.mime_type,
                "file_size": link.document.file_size,
                "created_at": link.document.created_at.isoformat(),
                "download_url": storage.get_download_url(link.document.storage_key)
            },
            "requirement": None,  # AI links are control-level, not requirement-level
            "note": "",
            "created_at": link.created_at.isoformat(),
            "confidence": link.confidence,
            "reasoning": link.reasoning,
            "is_ai_linked": True
        })

    # Get manually linked evidence (EvidenceLink)
    manual_links = db.query(EvidenceLink).filter(
        EvidenceLink.control_id == control.id
    ).all()

    for link in manual_links:
        result.append({
            "id": str(link.id),
            "document": {
                "id": str(link.document.id),
                "filename": link.document.filename,
                "mime_type": link.document.mime_type,
                "file_size": link.document.file_size,
                "created_at": link.document.created_at.isoformat(),
                "download_url": storage.get_download_url(link.document.storage_key)
            },
            "requirement": {
                "id": str(link.requirement.id),
                "req_code": link.requirement.req_code,
                "text": link.requirement.text
            } if link.requirement else None,
            "note": link.note,
            "created_at": link.created_at.isoformat(),
            "confidence": None,
            "reasoning": None,
            "is_ai_linked": False
        })

    return {"evidence": result}

@app.get("/controls/{control_id}/documents")
async def get_control_documents(control_id: str, db: Session = Depends(get_db)):
    """Get all documents linked to a specific control via AI analysis."""
    try:
        control = get_control_by_id_or_code(db, control_id)
        if not control:
            raise HTTPException(status_code=404, detail="Control not found")

        # Get all document links for this control
        links = db.query(DocumentControlLink).filter(DocumentControlLink.control_id == control.id).all()
        
        documents = []
        for link in links:
            document = link.document
            if document:
                documents.append({
                    "id": str(document.id),
                    "filename": document.filename,
                    "mime_type": document.mime_type,
                    "file_size": document.file_size,
                    "created_at": document.created_at.isoformat(),
                    "download_url": f"/api/documents/{document.id}/download",
                    "confidence": link.confidence,
                    "reasoning": link.reasoning,
                    "link_created_at": link.created_at.isoformat(),
                    "link_id": str(link.id),
                    "is_ai_linked": True
                })
        
        return {
            "control": {
                "id": str(control.id),
                "code": control.code,
                "title": control.title
            },
            "documents": documents
        }
        
    except Exception as e:
        logger.error(f"Failed to get documents for control {control_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to get control documents")

@app.delete("/document-control-links/{link_id}")
async def remove_document_control_link(link_id: str, db: Session = Depends(get_db)):
    """Remove an AI-generated document-control link (for false positives)."""
    try:
        link = db.query(DocumentControlLink).filter(DocumentControlLink.id == link_id).first()
        if not link:
            raise HTTPException(status_code=404, detail="Link not found")

        # Log the removal for audit
        logger.info(f"Removing AI document-control link: {link.document.filename} → Control {link.control.code} (confidence: {link.confidence})")

        db.delete(link)
        db.commit()

        return {"message": "Link removed successfully"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to remove document-control link {link_id}: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to remove link")

@app.post("/controls/{control_id}/scan")
async def create_scan(
    control_id: str,
    db: Session = Depends(get_db)
):
    """Create a new compliance scan for a control."""

    # Verify control exists
    control = get_control_by_id_or_code(db, control_id)
    if not control:
        raise HTTPException(status_code=404, detail="Control not found")

    # For demo, get default org
    org = db.query(Org).first()
    if not org:
        raise HTTPException(status_code=400, detail="No organization found")

    # Check if there's evidence linked to this control (manual OR AI-linked)
    manual_evidence_count = db.query(EvidenceLink).filter(
        EvidenceLink.control_id == control.id,
        EvidenceLink.org_id == org.id
    ).count()

    ai_evidence_count = db.query(DocumentControlLink).filter(
        DocumentControlLink.control_id == control.id
    ).count()

    total_evidence = manual_evidence_count + ai_evidence_count

    if total_evidence == 0:
        raise HTTPException(
            status_code=400,
            detail="No evidence linked to this control. Please upload and link evidence documents first."
        )

    # Create scan record
    scan = Scan(
        org_id=org.id,
        control_id=control.id,
        status='pending'
    )
    
    db.add(scan)
    db.commit()
    db.refresh(scan)
    
    # Start background scan processing
    process_scan.delay(str(scan.id))
    
    return {
        "scan_id": str(scan.id),
        "status": "pending",
        "message": "Scan started. Check status using the scan_id."
    }

@app.get("/scans/{scan_id}")
async def get_scan_status(scan_id: str, db: Session = Depends(get_db)):
    """Get the status and results of a scan."""
    
    scan = db.query(Scan).filter(Scan.id == scan_id).first()
    if not scan:
        raise HTTPException(status_code=404, detail="Scan not found")
    
    # Get scan results
    results = db.query(ScanResult).filter(ScanResult.scan_id == scan.id).all()
    gaps = db.query(Gap).filter(Gap.scan_id == scan.id).all()
    
    return {
        "id": str(scan.id),
        "control": {
            "id": str(scan.control.id),
            "code": scan.control.code,
            "title": scan.control.title
        },
        "status": scan.status,
        "model": scan.model,
        "prompt_version": scan.prompt_version,
        "progress_percentage": scan.progress_percentage or 0,
        "current_step": scan.current_step or 'Initializing...',
        "total_requirements": scan.total_requirements or 0,
        "processed_requirements": scan.processed_requirements or 0,
        "created_at": scan.created_at.isoformat(),
        "updated_at": scan.updated_at.isoformat(),
        "results": [
            {
                "requirement": {
                    "id": str(result.requirement.id),
                    "req_code": result.requirement.req_code,
                    "text": result.requirement.text,
                    "maturity_level": result.requirement.maturity_level
                },
                "outcome": result.outcome,
                "confidence": result.confidence,
                "rationale": _safe_json_loads(result.rationale_json),
                "citations": _safe_json_loads(result.citations_json, default=[])
            }
            for result in results
        ],
        "gaps": [
            {
                "requirement": {
                    "id": str(gap.requirement.id),
                    "req_code": gap.requirement.req_code,
                    "text": gap.requirement.text
                },
                "summary": gap.gap_summary,
                "recommended_actions": _safe_json_loads(gap.recommended_actions_json, default=[])
            }
            for gap in gaps
        ]
    }

@app.get("/controls/{control_id}/scans")
async def get_control_scans(control_id: str, db: Session = Depends(get_db)):
    """Get all scans for a control."""

    # Find the control by ID or code
    control = get_control_by_id_or_code(db, control_id)
    if not control:
        raise HTTPException(status_code=404, detail="Control not found")

    scans = db.query(Scan).filter(
        Scan.control_id == control.id
    ).order_by(Scan.created_at.desc()).all()
    
    return {
        "scans": [
            {
                "id": str(scan.id),
                "status": scan.status,
                "model": scan.model,
                "prompt_version": scan.prompt_version,
                "progress_percentage": scan.progress_percentage or 0,
                "current_step": scan.current_step or 'Initializing...',
                "total_requirements": scan.total_requirements or 0,
                "processed_requirements": scan.processed_requirements or 0,
                "created_at": scan.created_at.isoformat()
            }
            for scan in scans
        ]
    }

# AI Settings endpoints
class AISettingsRequest(BaseModel):
    provider: str  # 'openai' or 'ollama'
    openai_api_key: Optional[str] = None
    openai_model: Optional[str] = 'gpt-4o'
    openai_vision_model: Optional[str] = 'gpt-4o'
    openai_endpoint: Optional[str] = None
    ollama_endpoint: Optional[str] = 'http://localhost:11434'
    ollama_model: Optional[str] = 'qwen2.5:14b'
    ollama_vision_model: Optional[str] = 'qwen2-vl'
    ollama_context_size: Optional[int] = 131072
    min_confidence_threshold: Optional[float] = 0.90
    use_dual_vision_validation: Optional[bool] = False

@app.get("/settings/ai")
async def get_ai_settings(db: Session = Depends(get_db)):
    """Get current AI provider settings from database."""
    # Get or create settings record
    settings = db.query(Settings).filter(Settings.id == 1).first()

    if not settings:
        # Create default settings if none exist
        settings = Settings(
            id=1,
            ai_provider=os.getenv("AI_PROVIDER", "ollama"),
            openai_model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
            openai_endpoint=os.getenv("OPENAI_ENDPOINT"),
            ollama_endpoint=os.getenv("OLLAMA_ENDPOINT", "http://host.docker.internal:11434"),
            ollama_model=os.getenv("OLLAMA_MODEL", "qwen2.5:14b"),
            ollama_context_size=int(os.getenv("OLLAMA_CONTEXT_SIZE", "131072"))
        )
        # Set API key from env if available
        if os.getenv("OPENAI_API_KEY"):
            settings.openai_api_key = os.getenv("OPENAI_API_KEY")

        db.add(settings)
        db.commit()
        db.refresh(settings)

    return {
        "provider": settings.ai_provider,
        "openai_model": settings.openai_model,
        "openai_vision_model": settings.openai_vision_model or 'gpt-4o',
        "openai_endpoint": settings.openai_endpoint,
        "ollama_endpoint": settings.ollama_endpoint,
        "ollama_model": settings.ollama_model,
        "ollama_vision_model": settings.ollama_vision_model or 'qwen2-vl',
        "ollama_context_size": settings.ollama_context_size,
        "min_confidence_threshold": settings.min_confidence_threshold or 0.90,
        "use_dual_vision_validation": bool(settings.use_dual_vision_validation) if hasattr(settings, 'use_dual_vision_validation') else False,
        # Don't return API key for security
        "openai_api_key": "***" if settings.openai_api_key else None
    }

@app.post("/settings/ai")
async def save_ai_settings(settings_request: AISettingsRequest, db: Session = Depends(get_db)):
    """Save AI provider settings to database."""
    # Get or create settings record
    settings = db.query(Settings).filter(Settings.id == 1).first()

    if not settings:
        settings = Settings(id=1)
        db.add(settings)

    # Update settings
    settings.ai_provider = settings_request.provider
    settings.updated_at = datetime.utcnow()

    # Update confidence threshold (applies to all providers)
    if settings_request.min_confidence_threshold is not None:
        settings.min_confidence_threshold = settings_request.min_confidence_threshold

    # Update dual vision validation setting
    if settings_request.use_dual_vision_validation is not None:
        settings.use_dual_vision_validation = bool(settings_request.use_dual_vision_validation)

    # Update OpenAI settings (allow even when provider is Ollama, for dual validation)
    if settings_request.openai_api_key and settings_request.openai_api_key != "***":
        settings.openai_api_key = settings_request.openai_api_key
    if settings_request.openai_model:
        settings.openai_model = settings_request.openai_model
    if settings_request.openai_vision_model:
        settings.openai_vision_model = settings_request.openai_vision_model
    if settings_request.openai_endpoint is not None:
        settings.openai_endpoint = settings_request.openai_endpoint if settings_request.openai_endpoint else None

    # Update Ollama settings (allow even when provider is OpenAI, for dual validation)
    if settings_request.ollama_endpoint:
        settings.ollama_endpoint = settings_request.ollama_endpoint
    if settings_request.ollama_model:
        settings.ollama_model = settings_request.ollama_model
    if settings_request.ollama_vision_model:
        settings.ollama_vision_model = settings_request.ollama_vision_model
    if settings_request.ollama_context_size is not None:
        settings.ollama_context_size = settings_request.ollama_context_size

    db.commit()
    db.refresh(settings)

    return {"message": "Settings saved successfully"}

@app.post("/settings/ai/test")
async def test_ai_connection(settings: AISettingsRequest):
    """Test connection to the specified AI provider."""
    try:
        if settings.provider == "openai":
            from openai import OpenAI
            
            if not settings.openai_api_key or settings.openai_api_key == "***":
                api_key = os.getenv("OPENAI_API_KEY")
            else:
                api_key = settings.openai_api_key
            
            # Use custom endpoint if provided
            base_url = settings.openai_endpoint if settings.openai_endpoint else None
            
            # Only require API key if using default OpenAI endpoint
            if not api_key and not base_url:
                raise HTTPException(status_code=400, detail="OpenAI API key is required for default OpenAI endpoint")
            
            # Use placeholder key for custom endpoints that don't require authentication
            if not api_key and base_url:
                api_key = LOCAL_AI_PLACEHOLDER_KEY

            client = OpenAI(
                api_key=api_key,
                base_url=base_url
            )
            response = create_chat_completion_safe(
                client=client,
                model=settings.openai_model or "gpt-4o-mini",
                messages=[{"role": "user", "content": "Reply with exactly: 'OpenAI connection successful'"}],
                max_tokens=10
            )
            
            return {
                "status": "success",
                "test_response": response.choices[0].message.content,
                "model": settings.openai_model or "gpt-4o-mini"
            }
            
        elif settings.provider == "ollama":
            import requests
            
            endpoint = settings.ollama_endpoint or "http://localhost:11434"
            model = settings.ollama_model or "llama2"
            
            # Test Ollama connection
            response = requests.post(
                f"{endpoint}/api/generate",
                json={
                    "model": model,
                    "prompt": "Reply with exactly: 'Ollama connection successful'",
                    "stream": False
                },
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                return {
                    "status": "success",
                    "test_response": result.get("response", "Connection successful"),
                    "model": model
                }
            else:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Ollama returned status {response.status_code}: {response.text}"
                )
        else:
            raise HTTPException(status_code=400, detail="Invalid AI provider")
            
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Connection test failed: {str(e)}")

@app.get("/settings/ollama/models")
async def get_ollama_models(endpoint: str = "http://localhost:11434"):
    """Get list of available models from Ollama instance."""
    try:
        import requests
        
        # Get list of models from Ollama
        response = requests.get(f"{endpoint}/api/tags", timeout=10)
        
        if response.status_code != 200:
            raise HTTPException(
                status_code=400, 
                detail=f"Cannot connect to Ollama at {endpoint}. Make sure Ollama is running."
            )
        
        data = response.json()
        models = []
        
        for model in data.get('models', []):
            name = model.get('name', '')
            size = model.get('size', 0)
            modified = model.get('modified_at', '')
            
            # Format model info
            size_gb = round(size / (1024**3), 1) if size > 0 else 0
            models.append({
                'name': name,
                'display_name': name.split(':')[0],  # Remove tag for display
                'size_gb': size_gb,
                'modified_at': modified,
                'family': get_model_family(name)
            })
        
        # Sort by family and name
        models.sort(key=lambda x: (x['family'], x['name']))
        
        return {
            "models": models,
            "endpoint": endpoint,
            "total_models": len(models)
        }
        
    except requests.exceptions.RequestException as e:
        raise HTTPException(
            status_code=400, 
            detail=f"Failed to connect to Ollama: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching models: {str(e)}")

@app.get("/settings/ollama/models/demo")
async def get_demo_ollama_models():
    """Demo response showing what Ollama models look like when available."""
    return {
        "models": [
            {
                "name": "llama3:8b",
                "display_name": "llama3",
                "size_gb": 4.7,
                "modified_at": "2024-12-25T10:00:00Z",
                "family": "Llama"
            },
            {
                "name": "llama2:7b",
                "display_name": "llama2",
                "size_gb": 3.8,
                "modified_at": "2024-12-20T15:30:00Z",
                "family": "Llama"
            },
            {
                "name": "mistral:7b",
                "display_name": "mistral",
                "size_gb": 4.1,
                "modified_at": "2024-12-22T12:00:00Z",
                "family": "Mistral"
            },
            {
                "name": "codellama:7b",
                "display_name": "codellama",
                "size_gb": 3.8,
                "modified_at": "2024-12-18T09:15:00Z",
                "family": "Code Llama"
            }
        ],
        "endpoint": "http://localhost:11434",
        "total_models": 4,
        "demo": True
    }

@app.get("/settings/openai/models")
async def get_openai_models(endpoint: str = None, api_key: str = None):
    """Get list of available models from OpenAI or custom OpenAI-compatible endpoint."""
    try:
        from openai import OpenAI
        
        # Use provided endpoint or fall back to environment/default
        base_url = endpoint if endpoint else os.getenv("OPENAI_ENDPOINT")
        
        # Use provided API key or fall back to environment
        if not api_key or api_key == "***":
            api_key = os.getenv("OPENAI_API_KEY")
        
        # Only require API key if using default OpenAI endpoint
        if not api_key and not base_url:
            raise HTTPException(status_code=400, detail="OpenAI API key is required for default OpenAI endpoint")
        
        # Use placeholder key for custom endpoints that don't require authentication
        if not api_key and base_url:
            api_key = LOCAL_AI_PLACEHOLDER_KEY

        # Create client with custom endpoint if provided
        client = OpenAI(
            api_key=api_key,
            base_url=base_url
        )
        
        # Query the /v1/models endpoint
        try:
            logger.info(f"Querying models from endpoint: {base_url or 'https://api.openai.com/v1'}")
            models_response = client.models.list()
            logger.info(f"Models response type: {type(models_response)}")
            
            # Handle case where models_response or data is None
            if not models_response:
                raise HTTPException(status_code=500, detail="No response from models endpoint")
            
            if not hasattr(models_response, 'data'):
                logger.error(f"Models response missing 'data' attribute. Available attributes: {dir(models_response)}")
                raise HTTPException(status_code=500, detail="Invalid response format - missing 'data' field")
            
            if models_response.data is None:
                raise HTTPException(status_code=500, detail="Models endpoint returned null data")
            
            logger.info(f"Models data type: {type(models_response.data)}, length: {len(models_response.data) if models_response.data else 'N/A'}")
            
            models = []
            if models_response.data:
                for i, model in enumerate(models_response.data):
                    logger.info(f"Processing model {i}: {type(model)}")
                    if model and hasattr(model, 'id'):
                        models.append({
                            'id': model.id,
                            'name': model.id,  # For compatibility with frontend
                            'display_name': model.id,
                            'created': getattr(model, 'created', 0),
                            'owned_by': getattr(model, 'owned_by', 'unknown'),
                            'object': getattr(model, 'object', 'model')
                        })
                    else:
                        logger.warning(f"Skipping invalid model at index {i}: {model}")
                    
        except AttributeError as e:
            raise HTTPException(status_code=500, detail=f"Unexpected response format from models endpoint: {str(e)}")
        except Exception as e:
            logger.error(f"OpenAI client models.list() failed: {str(e)}")
            
            # Fallback: try direct HTTP request
            if base_url:
                try:
                    import requests
                    headers = {"Authorization": f"Bearer {api_key}"} if api_key != LOCAL_AI_PLACEHOLDER_KEY else {}
                    models_url = f"{base_url.rstrip('/')}/models"
                    logger.info(f"Trying direct HTTP request to: {models_url}")
                    
                    response = requests.get(models_url, headers=headers, timeout=10)
                    
                    if response.status_code == 200:
                        data = response.json()
                        logger.info(f"Direct HTTP response: {data}")
                        
                        models = []
                        model_list = data.get('data', []) if isinstance(data, dict) else []
                        
                        for model in model_list:
                            if isinstance(model, dict) and 'id' in model:
                                models.append({
                                    'id': model['id'],
                                    'name': model['id'],
                                    'display_name': model['id'],
                                    'created': model.get('created', 0),
                                    'owned_by': model.get('owned_by', 'unknown'),
                                    'object': model.get('object', 'model')
                                })
                        
                        if models:
                            models.sort(key=lambda x: x['name'])
                            return {
                                "models": models,
                                "endpoint": base_url,
                                "total_models": len(models)
                            }
                    
                    logger.error(f"Direct HTTP request failed: {response.status_code} - {response.text}")
                    
                except Exception as fallback_error:
                    logger.error(f"Fallback HTTP request also failed: {fallback_error}")
            
            if "models" in str(e).lower() or "404" in str(e):
                raise HTTPException(status_code=404, detail="Models endpoint not found - may not be OpenAI-compatible")
            else:
                raise HTTPException(status_code=500, detail=f"Error querying models: {str(e)}")
        
        # Sort by name
        models.sort(key=lambda x: x['name'])
        
        return {
            "models": models,
            "endpoint": base_url or "https://api.openai.com/v1",
            "total_models": len(models)
        }
        
    except Exception as e:
        error_msg = str(e).lower()
        if 'api key' in error_msg or 'authentication' in error_msg or 'unauthorized' in error_msg:
            raise HTTPException(status_code=401, detail="Invalid API key or authentication failed")
        elif 'not found' in error_msg or '404' in error_msg:
            raise HTTPException(status_code=404, detail="Models endpoint not found - may not be OpenAI-compatible")
        elif 'connection' in error_msg or 'timeout' in error_msg:
            raise HTTPException(status_code=503, detail="Cannot connect to the endpoint")
        else:
            raise HTTPException(status_code=500, detail=f"Failed to fetch models: {str(e)}")

def create_chat_completion_safe(client, model, messages, max_tokens=None, temperature=None, use_json_mode=False):
    """
    Create a chat completion with safe fallbacks for different endpoint capabilities.
    Some OpenAI-compatible endpoints don't support all features.
    """
    # Base parameters
    params = {
        "model": model,
        "messages": messages
    }
    
    # Add optional parameters
    if max_tokens is not None:
        params["max_tokens"] = max_tokens
    if temperature is not None:
        params["temperature"] = temperature
    
    # Try with JSON mode first if requested
    if use_json_mode:
        try:
            params["response_format"] = {"type": "json_object"}
            return client.chat.completions.create(**params)
        except Exception as e:
            # If JSON mode fails, try without it
            logger.warning(f"JSON mode not supported, falling back to text mode: {e}")
            params.pop("response_format", None)
    
    # Make request without JSON mode
    return client.chat.completions.create(**params)

def get_model_family(model_name: str) -> str:
    """Categorize model by family for better organization."""
    name_lower = model_name.lower()
    
    if 'llama' in name_lower:
        return 'Llama'
    elif 'mistral' in name_lower:
        return 'Mistral'
    elif 'mixtral' in name_lower:
        return 'Mixtral'
    elif 'codellama' in name_lower:
        return 'Code Llama'
    elif 'phi' in name_lower:
        return 'Phi'
    elif 'gemma' in name_lower:
        return 'Gemma'
    elif 'qwen' in name_lower:
        return 'Qwen'
    else:
        return 'Other'

# Additional API endpoints for control details and reporting

@app.get("/controls/{control_id}")
async def get_control_details(control_id: str):
    """Get detailed information about a specific control including requirements."""
    db = SessionLocal()
    try:
        control = get_control_by_id_or_code(db, control_id)
        if not control:
            raise HTTPException(status_code=404, detail="Control not found")
        
        requirements = db.query(Requirement).filter(Requirement.control_id == control.id).all()
        
        return {
            "id": str(control.id),
            "framework_id": str(control.framework_id),
            "framework_name": control.framework.name,
            "code": control.code,
            "title": control.title,
            "description": control.description,
            "requirements": [
                {
                    "id": str(req.id),
                    "req_code": req.req_code,
                    "text": req.text,
                    "maturity_level": req.maturity_level,
                    "guidance": req.guidance
                }
                for req in requirements
            ],
            "created_at": control.created_at.isoformat()
        }
    finally:
        db.close()

@app.get("/controls/{control_id}/scans")
async def get_control_scans(control_id: str):
    """Get all scans for a specific control."""
    db = SessionLocal()
    try:
        # Find the control by ID or code
        control = get_control_by_id_or_code(db, control_id)
        if not control:
            raise HTTPException(status_code=404, detail="Control not found")

        scans = db.query(Scan).filter(
            Scan.control_id == control.id
        ).order_by(Scan.created_at.desc()).all()
        
        scan_list = []
        for scan in scans:
            scan_list.append({
                "id": str(scan.id),
                "status": scan.status,
                "model": scan.model,
                "prompt_version": scan.prompt_version,
                "progress_percentage": scan.progress_percentage or 0,
                "current_step": scan.current_step or 'Initializing...',
                "total_requirements": scan.total_requirements or 0,
                "processed_requirements": scan.processed_requirements or 0,
                "created_at": scan.created_at.isoformat(),
                "updated_at": scan.updated_at.isoformat()
            })
        
        return {"scans": scan_list}
    finally:
        db.close()

# AI-powered document analysis endpoints
class ControlAnalysisRequest(BaseModel):
    prompt: str
    max_tokens: Optional[int] = 1000
    temperature: Optional[float] = 0.3

class DocumentBatchAnalysisRequest(BaseModel):
    documents: List[dict]  # [{filename, content, type}]
    controls: List[dict]   # [{code, title, framework, description, evidence_types}]
    prompt: str

@app.post("/ai/analyze-text")
async def analyze_text_with_ai(request: ControlAnalysisRequest):
    """Analyze text using the configured AI provider."""
    try:
        from ai_scanner import get_ai_client
        import json as json_module
        
        ai_client = get_ai_client()
        
        if isinstance(ai_client, dict) and ai_client.get('type') == 'ollama':
            # Handle Ollama
            import requests
            
            endpoint = ai_client['endpoint']
            model = ai_client['model']
            
            response = requests.post(
                f"{endpoint}/api/generate",
                json={
                    "model": model,
                    "prompt": request.prompt,
                    "stream": False,
                    "options": {
                        "temperature": request.temperature,
                        "num_predict": request.max_tokens,
                        "num_ctx": int(os.getenv("OLLAMA_CONTEXT_SIZE", "32768"))
                    }
                },
                timeout=60
            )
            
            if response.status_code != 200:
                raise HTTPException(
                    status_code=500, 
                    detail=f"Ollama API error: {response.status_code} - {response.text}"
                )
            
            result = response.json()
            ai_response = result.get('response', '')
            
            # Check thinking field if response is empty (some models use this field)
            if not ai_response and 'thinking' in result:
                ai_response = result.get('thinking', '')
            
        else:
            # Handle OpenAI
            model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
            
            response = create_chat_completion_safe(
                client=ai_client,
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are a helpful AI assistant that analyzes documents and provides structured responses."
                    },
                    {
                        "role": "user", 
                        "content": request.prompt
                    }
                ],
                max_tokens=request.max_tokens,
                temperature=request.temperature
            )
            
            ai_response = response.choices[0].message.content
        
        return {"response": ai_response}

    except Exception as e:
        logger.error(f"AI analysis failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"AI analysis failed: {str(e)}"
        )

@app.post("/api/ai/analyze-image")
async def analyze_image_with_ai(
    image: UploadFile = File(...),
    prompt: str = Form(...)
):
    """Analyze an image using vision AI and suggest compliance controls."""
    try:
        from ai_scanner import get_ai_client
        import base64
        from PIL import Image
        import io
        import json as json_module

        # Read image content
        image_content = await image.read()

        # Process image with PIL
        try:
            pil_image = Image.open(io.BytesIO(image_content))
            # Convert to RGB if needed
            if pil_image.mode != 'RGB':
                pil_image = pil_image.convert('RGB')

            # Resize if too large
            max_size = (1024, 1024)
            if pil_image.size[0] > max_size[0] or pil_image.size[1] > max_size[1]:
                pil_image.thumbnail(max_size, Image.Resampling.LANCZOS)

            # Convert to bytes
            img_buffer = io.BytesIO()
            pil_image.save(img_buffer, format='JPEG', quality=85)
            processed_content = img_buffer.getvalue()
        except Exception:
            processed_content = image_content

        ai_client = get_ai_client()

        if not isinstance(ai_client, dict):  # OpenAI
            try:
                image_b64 = base64.b64encode(processed_content).decode('utf-8')

                response = ai_client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": prompt
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{image_b64}",
                                        "detail": "high"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=2000
                )

                ai_response = response.choices[0].message.content
                logger.info(f"Vision AI analysis complete for {image.filename}")

                return {"response": ai_response}

            except Exception as e:
                logger.warning(f"Vision AI failed for {image.filename}: {e}")
                # Fallback to OCR
                import pytesseract
                ocr_text = pytesseract.image_to_string(Image.open(io.BytesIO(image_content)))

                if ocr_text.strip():
                    # Analyze OCR text with the prompt
                    response = ai_client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "user",
                                "content": f"{prompt}\n\nExtracted text from image:\n{ocr_text[:3000]}"
                            }
                        ],
                        max_tokens=2000
                    )
                    ai_response = response.choices[0].message.content
                    return {"response": ai_response}
                else:
                    raise HTTPException(
                        status_code=500,
                        detail="Could not extract text from image"
                    )
        else:  # Ollama
            # Use OCR for Ollama
            import pytesseract
            ocr_text = pytesseract.image_to_string(Image.open(io.BytesIO(image_content)))

            if not ocr_text.strip():
                raise HTTPException(
                    status_code=500,
                    detail="Could not extract text from image with OCR"
                )

            # Analyze OCR text with Ollama
            import requests
            endpoint = ai_client['endpoint']
            model = ai_client['model']

            analysis_prompt = f"{prompt}\n\nExtracted text from image:\n{ocr_text[:3000]}"

            response = requests.post(
                f"{endpoint}/api/generate",
                json={
                    "model": model,
                    "prompt": analysis_prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.3,
                        "num_predict": 2000,
                        "num_ctx": int(os.getenv("OLLAMA_CONTEXT_SIZE", "32768"))
                    }
                },
                timeout=120
            )

            if response.status_code != 200:
                raise HTTPException(
                    status_code=500,
                    detail=f"Ollama API error: {response.status_code}"
                )

            result = response.json()
            ai_response = result.get('response', '')

            if not ai_response and 'thinking' in result:
                ai_response = result.get('thinking', '')

            logger.info(f"OCR + Ollama analysis complete for {image.filename}")
            return {"response": ai_response}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Image analysis failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Image analysis failed: {str(e)}"
        )

@app.post("/analyze-documents")
async def analyze_multiple_documents(request: DocumentBatchAnalysisRequest):
    """Analyze multiple documents together and suggest relevant compliance controls."""
    try:
        from ai_scanner import get_ai_client
        import json as json_module
        
        ai_client = get_ai_client()
        
        # Prepare the analysis prompt with all document information
        documents_summary = []
        for doc in request.documents:
            doc_info = f"Document: {doc['filename']} ({doc['type']})"
            if doc.get('content') and doc['content'].strip():
                doc_info += f"\nContent preview: {doc['content'][:8000]}{'...' if len(doc['content']) > 8000 else ''}"
            documents_summary.append(doc_info)
        
        controls_summary = []
        for ctrl in request.controls:
            ctrl_info = f"{ctrl['code']}: {ctrl['title']} ({ctrl['framework']})"
            if ctrl.get('evidence_types'):
                ctrl_info += f" - Evidence needed: {ctrl['evidence_types']}"
            controls_summary.append(ctrl_info)
        
        enhanced_prompt = f"""
{request.prompt}

Documents to analyze:
{chr(10).join(documents_summary)}

Available compliance controls:
{chr(10).join(controls_summary)}

IMPORTANT: You must respond with ONLY valid JSON in this exact format:
{{
  "suggestions": [
    {{
      "control_code": "CONTROL_CODE",
      "control_title": "Control Title",
      "framework_name": "Framework Name", 
      "confidence": 0.8,
      "reasoning": "Brief explanation mentioning specific documents"
    }}
  ]
}}

Do not include any text before or after the JSON. Do not use markdown formatting. Return only the JSON object."""

        result = None
        
        if isinstance(ai_client, dict) and ai_client.get('type') == 'ollama':
            # Handle Ollama
            import requests
            
            endpoint = ai_client['endpoint']
            model = ai_client['model']
            
            response = requests.post(
                f"{endpoint}/api/generate",
                json={
                    "model": model,
                    "prompt": enhanced_prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.3,
                        "num_predict": 2000,
                        "num_ctx": int(os.getenv("OLLAMA_CONTEXT_SIZE", "32768"))
                    }
                },
                timeout=45  # Reduced timeout to prevent connection drops
            )
            
            if response.status_code == 200:
                result_json = response.json()
                result = result_json.get("response", "")
                
                # Check thinking field if response is empty (some models use this field)
                if not result and 'thinking' in result_json:
                    result = result_json.get('thinking', '')
            else:
                raise HTTPException(status_code=500, detail=f"Ollama API error: {response.text}")
                
        else:
            # Handle OpenAI
            from openai import OpenAI
            client = OpenAI(api_key=ai_client.api_key, base_url=ai_client.base_url)
            
            completion = client.chat.completions.create(
                model=ai_client.model,
                messages=[
                    {"role": "system", "content": "You are a compliance expert. Analyze documents and suggest relevant compliance controls in JSON format."},
                    {"role": "user", "content": enhanced_prompt}
                ],
                max_tokens=2000,
                temperature=0.3,
                response_format={"type": "json_object"}
            )
            
            result = completion.choices[0].message.content
        
        # Try to parse the response as JSON
        try:
            if result:
                parsed_result = _extract_json_from_response(result)
                if parsed_result:
                    return {"suggestions": parsed_result.get("suggestions", [])}
                else:
                    logger.warning(f"Could not extract JSON from batch analysis response: {result[:200]}...")
                    return {"suggestions": []}
            else:
                return {"suggestions": []}
        except Exception as e:
            # If JSON parsing fails, return the raw result
            logger.warning(f"Error parsing batch analysis response: {e}")
            return {"suggestions": [], "raw_response": result}
            
    except Exception as e:
        logger.error(f"Batch document analysis failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Batch analysis failed: {str(e)}"
        )

@app.post("/analyze-document-controls")
async def analyze_document_controls(
    file: UploadFile = File(...),
    available_controls: str = None
):
    """Analyze uploaded document and suggest relevant compliance controls."""
    try:
        # Read file content
        file_content = await file.read()
        file_text = ""
        
        # Extract text based on file type
        if file.content_type == "text/plain":
            file_text = file_content.decode('utf-8')
        elif file.content_type == "application/pdf":
            # Extract text from PDF
            try:
                import fitz  # PyMuPDF
                import io
                
                pdf_doc = fitz.open(stream=file_content, filetype="pdf")
                text_pages = []
                for page_num in range(min(3, pdf_doc.page_count)):  # First 3 pages
                    page = pdf_doc[page_num]
                    page_text = page.get_text()
                    if page_text.strip():
                        text_pages.append(f"Page {page_num + 1}: {page_text[:1000]}")
                
                pdf_doc.close()
                file_text = "\n\n".join(text_pages)
                if not file_text.strip():
                    file_text = f"PDF document: {file.filename} (text extraction failed)"
            except Exception as e:
                logger.warning(f"PDF text extraction failed: {e}")
                file_text = f"PDF document: {file.filename} (text extraction failed)"
        
        elif file.content_type.startswith("image/"):
            # Analyze image using AI vision
            try:
                import base64
                
                # Convert image to base64 for AI analysis
                image_b64 = base64.b64encode(file_content).decode('utf-8')
                
                # Use AI to describe the image content
                from ai_scanner import get_ai_client
                ai_client = get_ai_client()
                
                if isinstance(ai_client, dict) and ai_client.get('type') == 'ollama':
                    # Ollama with vision models (if available)
                    try:
                        import requests
                        endpoint = ai_client['endpoint']
                        
                        # Try vision model first
                        vision_response = requests.post(
                            f"{endpoint}/api/generate",
                            json={
                                "model": "llava",  # Vision model
                                "prompt": f"Describe what you see in this image. Focus on any text, security-related content, error messages, configurations, or compliance-related information: {file.filename}",
                                "images": [image_b64],
                                "stream": False
                            },
                            timeout=30
                        )
                        
                        if vision_response.status_code == 200:
                            result = vision_response.json()
                            file_text = f"Image analysis of {file.filename}: {result.get('response', '')}"
                        else:
                            raise Exception("Vision model not available")
                    except:
                        # Fallback to filename analysis
                        file_text = f"Image: {file.filename} (visual analysis not available)"
                        
                else:
                    # OpenAI GPT-4 Vision
                    try:
                        response = ai_client.chat.completions.create(
                            model="gpt-4-vision-preview",
                            messages=[
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": f"Analyze this image and describe any text, security configurations, error messages, compliance-related information, or other relevant content you can see. Image filename: {file.filename}"
                                        },
                                        {
                                            "type": "image_url",
                                            "image_url": {
                                                "url": f"data:{file.content_type};base64,{image_b64}"
                                            }
                                        }
                                    ]
                                }
                            ],
                            max_tokens=500
                        )
                        file_text = f"Image analysis of {file.filename}: {response.choices[0].message.content}"
                    except Exception as e:
                        logger.warning(f"Vision analysis failed: {e}")
                        file_text = f"Image: {file.filename} (visual analysis failed)"
                        
            except Exception as e:
                logger.error(f"Image analysis error: {e}")
                file_text = f"Image: {file.filename}"
                
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Extract text from Word documents
            try:
                import docx
                import io
                
                doc = docx.Document(io.BytesIO(file_content))
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                file_text = "\n".join(paragraphs[:20])  # First 20 paragraphs
                if not file_text.strip():
                    file_text = f"Word document: {file.filename} (text extraction failed)"
            except Exception as e:
                logger.warning(f"Word document text extraction failed: {e}")
                file_text = f"Word document: {file.filename} (text extraction failed)"
                
        else:
            # For other file types, use filename
            file_text = f"Document: {file.filename}"
        
        # Parse available controls
        controls = []
        if available_controls:
            try:
                controls = json_module.loads(available_controls)
            except:
                pass
        
        if not controls:
            return {"suggested_controls": []}
        
        # Create analysis prompt
        controls_context = "\n".join([
            f"- {c.get('code', 'N/A')}: {c.get('title', 'N/A')} ({c.get('framework', 'N/A')}) - {c.get('description', 'N/A')[:500]}..."
            for c in controls[:50]  # Increased from 10 to 50 controls with larger context
        ])
        
        analysis_prompt = f"""
Analyze this document and determine which compliance controls it might relate to:

Document: {file.filename}
Content preview: {file_text[:5000]}{'...' if len(file_text) > 5000 else ''}

Available compliance controls:
{controls_context}

BE EXTREMELY STRICT with confidence scores. Use these guidelines:
- 0.8-1.0: ONLY for explicit, comprehensive policy documents with clear compliance statements
- 0.6-0.7: Strong documentation with specific compliance details
- 0.4-0.5: Partial evidence or screenshots with limited context
- 0.2-0.3: Weak evidence, filename-only matching, or requires significant interpretation
- 0.0-0.1: No clear relevance

CRITICAL: Screenshots or images alone should receive LOW confidence (0.2-0.4) unless they show comprehensive,
unambiguous compliance with clear context.

Respond in JSON format with a "suggestions" array containing objects with:
- control_code: the control code
- control_title: the control title
- framework_name: the framework name
- confidence: score from 0.0 to 1.0 (BE STRICT - most should be < 0.5)
- reasoning: brief explanation

Limit to the top 3 most relevant matches. If no relevant matches, return empty array.
"""
        
        # Call AI analysis
        from ai_scanner import get_ai_client
        
        ai_client = get_ai_client()
        
        if isinstance(ai_client, dict) and ai_client.get('type') == 'ollama':
            # Handle Ollama
            import requests
            
            endpoint = ai_client['endpoint']
            model = ai_client['model']
            
            response = requests.post(
                f"{endpoint}/api/generate",
                json={
                    "model": model,
                    "prompt": analysis_prompt + """\n\nIMPORTANT: You must respond with ONLY valid JSON in this exact format:
{{
  "suggestions": [
    {{
      "control_code": "CONTROL_CODE",
      "control_title": "Control Title",
      "framework_name": "Framework Name", 
      "confidence": 0.8,
      "reasoning": "Brief explanation"
    }}
  ]
}}

Do not include any text before or after the JSON. Do not use markdown formatting. Return only the JSON object.""",
                    "stream": False,
                    "options": {
                        "temperature": 0.3,
                        "num_predict": 2000,
                        "num_ctx": int(os.getenv("OLLAMA_CONTEXT_SIZE", "32768"))
                    }
                },
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result.get('response', '')
                
                # Check thinking field if response is empty (some models use this field)
                if not ai_response and 'thinking' in result:
                    ai_response = result.get('thinking', '')
            else:
                raise Exception(f"Ollama error: {response.status_code}")
                
        else:
            # Handle OpenAI
            model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
            
            response = create_chat_completion_safe(
                client=ai_client,
                model=model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are a compliance expert that analyzes documents and suggests relevant compliance controls. Respond only with valid JSON."
                    },
                    {
                        "role": "user", 
                        "content": analysis_prompt
                    }
                ],
                max_tokens=800,
                temperature=0.3,
                use_json_mode=True
            )
            
            ai_response = response.choices[0].message.content
        
        # Parse AI response
        try:
            # Use improved JSON extraction
            parsed_response = _extract_json_from_response(ai_response)
            if not parsed_response:
                logger.warning(f"Could not extract JSON from AI response: {ai_response[:200]}...")
                return {"suggested_controls": []}
            
            suggestions = parsed_response.get('suggestions', [])
            
            # Validate and clean suggestions
            valid_suggestions = []
            for suggestion in suggestions[:3]:  # Limit to top 3
                if all(key in suggestion for key in ['control_code', 'control_title', 'confidence', 'reasoning']):
                    # Ensure confidence is between 0 and 1
                    confidence = max(0.0, min(1.0, float(suggestion.get('confidence', 0))))
                    valid_suggestions.append({
                        'control_code': str(suggestion['control_code']),
                        'control_title': str(suggestion['control_title']),
                        'framework_name': str(suggestion.get('framework_name', 'Unknown')),
                        'confidence': confidence,
                        'reasoning': str(suggestion['reasoning'])
                    })
            
            return {"suggested_controls": valid_suggestions}
            
        except json_module.JSONDecodeError:
            logger.warning(f"Failed to parse AI response as JSON: {ai_response[:200]}...")
            return {"suggested_controls": []}
        
    except Exception as e:
        logger.error(f"Document control analysis failed: {e}")
        # Return empty suggestions on error rather than failing  
        return {"suggested_controls": []}
        
    finally:
        # Reset file position for any subsequent reads
        await file.seek(0)